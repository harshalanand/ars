"""Generate ARS Listing Process SOP as a DOCX document."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

for section in doc.sections:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.8); section.right_margin = Cm(1.8)

style = doc.styles["Normal"]; style.font.name = "Calibri"; style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4); style.paragraph_format.space_before = Pt(2)
for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]; hs.font.color.rgb = RGBColor(0x1B,0x3A,0x5C); hs.font.name = "Calibri"

def T(doc, headers, rows, cw=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
        c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>'))
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c=t.rows[ri+1].cells[ci]; c.text=str(val)
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(9)
            if ri%2==1: c._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="EDF2F7"/>'))
    if cw:
        for i,w in enumerate(cw):
            for row in t.rows: row.cells[i].width=Cm(w)
    return t

def C(doc, text):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run(text); r.font.name="Consolas"; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0x1A,0x1A,0x2E)
    p._p.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F0F0F0"/>'))

def B(p, text): r=p.add_run(text); r.bold=True; return r

def BL(doc, text):
    p=doc.add_paragraph(text, style="List Bullet"); p.paragraph_format.left_indent=Cm(1.2)
    for r in p.runs: r.font.size=Pt(10)
    return p


# ══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════
doc.add_paragraph(); doc.add_paragraph()
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("ARS Listing Process"); r.font.size=Pt(28); r.font.color.rgb=RGBColor(0x1B,0x3A,0x5C); r.bold=True
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("Step-by-Step SOP & Technical Review"); r.font.size=Pt(16); r.font.color.rgb=RGBColor(0x4A,0x6E,0x8C)
doc.add_paragraph()
m=doc.add_paragraph(); m.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=m.add_run("V2 Retail Auto Replenishment System\nVersion 2.0 — April 2026\n\nOwner: Akash Agarwal, Director V2 Retail\nRepository: github.com/harshalanand/ars")
r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x55,0x55,0x55)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading("Table of Contents", level=1)
for item in [
    "1. Overview & Terminology",
    "2. Output Tables",
    "3. Phase 0: Configuration & Settings",
    "4. Phase 1: Pre-Grid Calculations (Full Pipeline)",
    "5. Phase 2: Listing Generation (Core)",
    "   5.1 RDC Mode Logic",
    "   5.2 Part 1: Grid Data (IS_NEW=0)",
    "   5.3 Part 2: MSA Missing Options (IS_NEW=1)",
    "   5.4 Part 2.5: Index Creation",
    "   5.5 Part 3.5: Data Enrichment (DPN, SAL_D, AGE, Sale)",
    "   5.6 Part 3.5a: Enrichment (LISTING, I_ROD, CLR, FOCUS)",
    "   5.7 Part 3.55: MSA Quantities + Variant Counts",
    "   5.8 Part 3.6: OPT_TYPE Classification (4-Way Tagging)",
    "   5.9 Part 3.7: MIX Aggregation",
    "   5.10 Part 4: Grid Joins (CONT, MBQ, OPT_CNT, DISP_Q)",
    "   5.11 Part 4b: PER_OPT_SALE",
    "   5.12 Part 4c: MANUAL_DENSITY DPN Override + OPT_MBQ",
    "   5.13 Part 4d: ART_EXCESS + Per-Grid REQ",
    "6. Phase 3: Store Ranking",
    "7. Phase 4: Working Table (ARS_LISTING_WORKING)",
    "8. Phase 5: Multi-Level Allocation (ARS_ALLOC_WORKING)",
    "   8.1 Create Alloc Table + Enrichment",
    "   8.2 Pool Tracker",
    "   8.3 Primary Allocation (RL > TBC > TBL, I_ROD Rounds)",
    "   8.4 Per-OPT Waterfall + Validation",
    "   8.5 Fallback (Grid Demotion)",
    "   8.6 Reflect to Working Table",
    "9. API Endpoints",
    "10. Data Flow Diagram",
]:
    p=doc.add_paragraph(item); p.paragraph_format.space_after=Pt(1)
    for r in p.runs: r.font.size=Pt(10)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. OVERVIEW
# ══════════════════════════════════════════════════════════════════
doc.add_heading("1. Overview & Terminology", level=1)
doc.add_paragraph(
    "The Listing module is the core data preparation engine of ARS. It combines existing store stock "
    "(from Grid tables) and MSA recommendations (what to replenish) into a unified master listing with "
    "calculated quantities, store rankings, and size-level allocation."
)
p=doc.add_paragraph(); B(p,"OPT "); p.add_run("(Option) = MAJ_CAT x GEN_ART_NUMBER x CLR. This is the fundamental unit throughout the listing process.")
p=doc.add_paragraph(); B(p,"Entry Point: "); p.add_run("POST /listing/generate")
p=doc.add_paragraph(); B(p,"Source: "); p.add_run("listing.py (~1,800 lines) + listing_allocator.py (~400 lines)")

doc.add_heading("Key Terminology", level=2)
T(doc, ["Term", "Meaning"], [
    ["OPT", "MAJ_CAT x GEN_ART_NUMBER x CLR (option = generated article + color)"],
    ["STK_TTL", "Total Stock = SUM of SLOC stock columns (excludes sale columns like L-7 DAYS SALE-Q)"],
    ["DPN", "Display Norm (target stock level per store x MAJ_CAT)"],
    ["SAL_D", "Sale Days (coverage period = INT_DAYS + PRD_DAYS + SL_CVR)"],
    ["OPT_MBQ", "Option MBQ = DPN + rate x SAL_D (target replenishment qty)"],
    ["FNL_Q", "Final Quantity from MSA = MAX(warehouse STK - PEND, 0)"],
    ["I_ROD", "Replenishment rounds (I_ROD=3 means OPT gets 3 allocation passes)"],
    ["MANUAL_DENSITY", "Article-level DPN override (if >0, replaces DPN in OPT_MBQ)"],
    ["CONT", "Contribution factor (0-1) — distributes qty to finer hierarchy/size levels"],
    ["ST_RANK", "Store priority rank per MAJ_CAT (1=best, gets stock first)"],
    ["IS_NEW", "0=existing store stock (from grid), 1=new MSA recommendation"],
], cw=[3, 14.5])

# ══════════════════════════════════════════════════════════════════
# 2. OUTPUT TABLES
# ══════════════════════════════════════════════════════════════════
doc.add_heading("2. Output Tables", level=1)
T(doc, ["Table", "Grain", "Purpose"], [
    ["ARS_LISTING", "Store x OPT", "Master listing — all options (existing + new)"],
    ["ARS_STORE_RANKING", "Store x MAJ_CAT", "Store prioritization per category"],
    ["ARS_LISTING_WORKING", "Store x OPT (filtered)", "Clean extract for allocation (eligible OPTs only)"],
    ["ARS_ALLOC_WORKING", "Store x OPT x VAR_ART x SZ", "Size-level allocation quantities (multi-level waterfall)"],
], cw=[4.5, 5.5, 7.5])

# ══════════════════════════════════════════════════════════════════
# 3. CONFIGURATION
# ══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("3. Phase 0: Configuration & Settings", level=1)
doc.add_paragraph("GET /listing/config loads available options. Settings auto-saved to AppSettings on each run.")
T(doc, ["Variable", "Default", "Purpose"], [
    ["stock_threshold_pct", "0.6 (60%)", "OPT_TYPE: RL when STK >= X% of DPN; also size availability threshold"],
    ["excess_multiplier", "2.0", "Excess: STK > X x OPT_MBQ = excess stock"],
    ["hold_days", "0", "Extra days added to SAL_D for IS_NEW=1 OPTs only"],
    ["age_threshold", "15", "OPTs with effective AGE < X use PER_OPT_SALE (IS_NEW=1 always 0)"],
    ["mix_mode", "st_maj_rng", "MIX aggregation: always at Store x MAJ_CAT (each = no agg)"],
    ["rdc_mode", "all", "RDC filter: all / own / cross"],
    ["req_weight", "0.4", "Store ranking: weight for requirement rank"],
    ["fill_weight", "0.6", "Store ranking: weight for fill rate rank"],
    ["enable_fallback", "false", "Enable fallback allocation (demote primary grids one by one)"],
], cw=[3.5, 2.5, 11.5])

# ══════════════════════════════════════════════════════════════════
# 4. PRE-GRID CALCULATIONS
# ══════════════════════════════════════════════════════════════════
doc.add_heading("4. Phase 1: Pre-Grid Calculations (Full Pipeline)", level=1)
doc.add_paragraph('Triggered when run_mode = "full". Builds ARS_CALC_ST_MAJ_CAT and ARS_CALC_ST_ART.')

doc.add_heading("SQL Migration (runs first)", level=2)
BL(doc, "DROP MANUAL_MBQ from Master_ALC_INPUT_CO_MAJ_CAT and ST_MAJ_CAT (not used at MAJ_CAT level)")
BL(doc, "RENAME MANUAL_MBQ to MANUAL_DENSITY in MASTER_ALC_INPUT_CO_ART and Master_ALC_INPUT_ST_ART")

doc.add_heading("MAJ_CAT Level (ARS_CALC_ST_MAJ_CAT)", level=2)
T(doc, ["Step", "What It Does"], [
    ["1. Create calc (CO base)", "CO_MAJ_CAT x all stores -> base table. Fallback: copy ST_MAJ_CAT."],
    ["1b. Fill CO gaps", "Insert missing (store x MAJ_CAT) combos from CO x ST_MASTER."],
    ["2. Overlay ST values", "ST_MAJ_CAT overrides CO per-store (non-null ST wins)."],
    ["3. Apply defaults", "LISTING->1, I_ROD->1, growth rates->1."],
    ["4. SAL_D", "INT_DAYS + PRD_DAYS + SL_CVR. Priority: ST > CO > ST_MASTER."],
    ["5. SAL_PD", "Per-day sale from CM/NM sale quantities and remaining days."],
], cw=[3.5, 14])

doc.add_heading("Article Level (ARS_CALC_ST_ART)", level=2)
T(doc, ["Step", "What It Does"], [
    ["A1. Create (CO base)", "CO_ART x stores -> base. Fallback: ST_ART. Drops CORE, AUTO, HH_ART."],
    ["A2. Overlay ST_ART", "ST_ART overrides CO values (same cascade as MAJ_CAT)."],
    ["A3. Defaults", "Same + FOCUS_W_CAP/FOCUS_WO_CAP: Y->1, else->0."],
    ["A4. MANUAL_DENSITY default", "<=0/null -> 0."],
    ["A5. DPN override", "If MANUAL_DENSITY > 0 -> DPN = MANUAL_DENSITY (article-level only)."],
    ["A6. SAL_D", "From ARS_CALC_ST_MAJ_CAT (article has no own SL_CVR)."],
    ["A7. SAL_PD", "From MASTER_GEN_ART_SALE + ST_MAJ for REM_D."],
], cw=[3.5, 14])

doc.add_heading("Master Sale (MASTER_GEN_ART_SALE)", level=2)
BL(doc, "Ensure MAJ_CAT column (from vw_master_product if missing)")
BL(doc, "Compute SAL_PD in-place on ~21L rows (feeds listing AUTO_GEN_ART_SALE)")

# ══════════════════════════════════════════════════════════════════
# 5. LISTING GENERATION
# ══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("5. Phase 2: Listing Generation (Core Pipeline)", level=1)
doc.add_paragraph("Builds ARS_LISTING. Table is dropped and recreated on each run.")

doc.add_heading("5.1 RDC Mode Logic", level=2)
T(doc, ["Mode", "Stores", "MSA Options"], [
    ["All", "All active stores (LISTING != 0/N)", "All MSA options, no RDC filter"],
    ["Own RDC", "Stores tagged to selected RDC(s)", "MSA options from same RDC(s) + RDC join"],
    ["Cross RDC", "Stores from cross_to RDC(s)", "MSA options from cross_from RDC(s)"],
], cw=[2.5, 6.5, 8.5])

doc.add_heading("5.2 Part 1: Grid Data (IS_NEW = 0)", level=2)
doc.add_paragraph("Insert existing store stock from ARS_GRID_MJ_GEN_ART. STK_TTL = SUM of stock SLOC columns only (sale columns like L-7 DAYS SALE-Q are excluded from sum but carried as columns).")

doc.add_heading("5.3 Part 2: MSA Missing Options (IS_NEW = 1)", level=2)
doc.add_paragraph("Insert MSA-recommended OPTs that the store does NOT have. Stock = 0, IS_NEW = 1.")

doc.add_heading("5.4 Part 2.5: Index Creation", level=2)
doc.add_paragraph("Only if listing >= 5,000 rows. Creates NONCLUSTERED INDEX on (WERKS, MAJ_CAT) and (GEN_ART_NUMBER).")

doc.add_heading("5.5 Part 3.5: Data Enrichment", level=2)
T(doc, ["Column", "Source", "Join Key"], [
    ["DPN", "ARS_CALC_ST_MAJ_CAT", "WERKS=ST_CD + MAJ_CAT"],
    ["SAL_D", "ARS_CALC_ST_MAJ_CAT", "WERKS=ST_CD + MAJ_CAT"],
    ["AUTO_GEN_ART_SALE", "MASTER_GEN_ART_SALE.SAL_PD", "WERKS + MAJ_CAT + GEN_ART + CLR"],
    ["AGE", "MASTER_GEN_ART_AGE", "WERKS + MAJ_CAT + GEN_ART + CLR"],
], cw=[4, 5.5, 8])

doc.add_heading("5.6 Part 3.5a: Enrichment (LISTING, I_ROD, CLR, FOCUS)", level=2)
doc.add_paragraph("Two-step cascade: ARS_CALC_ST_MAJ_CAT first, then ARS_CALC_ST_ART overrides.")
T(doc, ["Column", "ST_MAJ_CAT (Step 1)", "ST_ART Cascade (Step 2)"], [
    ["LISTING", "Yes", "Yes (override where non-null/non-zero)"],
    ["I_ROD", "Yes", "Yes (override where non-null/non-zero)"],
    ["CLR_MIN", "Yes", "No (MAJ_CAT level only)"],
    ["CLR_MAX", "Yes", "No (MAJ_CAT level only)"],
    ["FOCUS_W_CAP", "No", "Yes (article-level only)"],
    ["FOCUS_WO_CAP", "No", "Yes (article-level only)"],
], cw=[3, 5, 9.5])
doc.add_paragraph("Note: MANUAL_DENSITY is NOT enriched here. It is used for DPN override in Part 4c.")

doc.add_heading("5.7 Part 3.55: MSA Quantities + Variant Counts", level=2)
T(doc, ["Column", "Source", "Logic"], [
    ["MSA_FNL_Q", "ARS_MSA_GEN_ART", "SUM(FNL_Q) per OPT [+ RDC in own mode]"],
    ["VAR_COUNT", "ARS_MSA_VAR_ART", "COUNT(*) of variant articles per OPT"],
    ["VAR_FNL_COUNT", "ARS_MSA_VAR_ART", "COUNT where FNL_Q > 0"],
], cw=[3, 4, 10.5])

doc.add_page_break()
doc.add_heading("5.8 Part 3.6: OPT_TYPE Classification (4-Way Tagging)", level=2)
doc.add_paragraph("Evaluated top-to-bottom, first match wins. Applies to ALL rows (IS_NEW=0 and IS_NEW=1).")
T(doc, ["OPT_TYPE", "Rule", "Meaning"], [
    ["MIX (a)", "STK < 60% x DPN AND MSA_FNL_Q = 0", "Low stock + no MSA backup"],
    ["MIX (b)", "VAR_FNL_COUNT / VAR_COUNT < 60%", "Poor size availability (all rows)"],
    ["RL", "STK >= 60% x DPN", "Adequate stock (Running Line)"],
    ["TBC", "0 < STK < 60% x DPN AND MSA_FNL_Q > 0", "Low stock + MSA available (To Be Checked)"],
    ["TBL", "STK <= 0 AND MSA_FNL_Q > 0", "Zero stock + MSA available (To Be Listed)"],
], cw=[2, 6, 9.5])

doc.add_heading("5.9 Part 3.7: MIX Aggregation", level=2)
doc.add_paragraph("ALL MIX-tagged rows (IS_NEW=0 and IS_NEW=1) aggregated into exactly 1 MIX row per Store x MAJ_CAT. Numeric cols summed; DPN/SAL_D fetched fresh from ARS_CALC_ST_MAJ_CAT.")

doc.add_heading("5.10 Part 4: Grid Joins", level=2)
doc.add_paragraph("Pre-resolve vw_master_product cols onto listing ONCE. Then join each active grid -> prefixed columns (MJ_STK_TTL, MJ_CONT, RNG_SEG_MBQ, etc.).")

doc.add_heading("5.11 Part 4b: PER_OPT_SALE", level=2)
C(doc, "PER_OPT_SALE = ((Grid_MBQ - Grid_DISP_Q) / Grid_DISP_Q x DPN) / SAL_D\nFrom the grid flagged use_for_opt_sale = 1.")

doc.add_heading("5.12 Part 4c: MANUAL_DENSITY DPN Override + OPT_MBQ", level=2)
doc.add_heading("DPN Override (before OPT_MBQ)", level=3)
C(doc,
    "Cascade (first match wins):\n"
    "  Source 1: ARS_CALC_ST_ART.MANUAL_DENSITY\n"
    "  Source 2: Master_ALC_INPUT_ST_ART.MANUAL_DENSITY (fallback)\n"
    "  Legacy:   MANUAL_MBQ column name also detected\n\n"
    "SQL: UPDATE ARS_LISTING SET DPN = MANUAL_DENSITY\n"
    "     WHERE MANUAL_DENSITY > 0\n"
    "     JOIN on WERKS=ST_CD + MAJ_CAT + GEN_ART_NUMBER [+ CLR]")

doc.add_heading("Effective AGE", level=3)
C(doc,
    "IS_NEW = 1 (any AGE)     -> 0   (always new)\n"
    "IS_NEW = 0, AGE = NULL   -> 0   (unknown = new)\n"
    "IS_NEW = 0, AGE = 0      -> 0   (new)\n"
    "IS_NEW = 0, AGE = 25     -> 25  (existing)")

doc.add_heading("Sale Rate Selection", level=3)
T(doc, ["Condition", "Rate"], [
    ["Effective AGE < 15", "MAX(PER_OPT_SALE, L-7/7, AUTO_GEN_ART_SALE)"],
    ["Effective AGE >= 15", "MAX(L-7/7, AUTO_GEN_ART_SALE)"],
], cw=[5, 12.5])

doc.add_heading("OPT_MBQ Formulas", level=3)
C(doc,
    "OPT_MBQ    = DPN + rate x SAL_D\n"
    "OPT_REQ    = MAX(0, OPT_MBQ - STK_TTL)\n"
    "OPT_MBQ_WH = DPN + rate x (SAL_D + HOLD_DAYS)  <- HOLD only for IS_NEW=1\n"
    "OPT_REQ_WH = MAX(0, OPT_MBQ_WH - STK_TTL)\n"
    "MAX_DAILY_SALE = MAX(L-7/7, AUTO_GEN_ART_SALE)")

doc.add_heading("5.13 Part 4d: ART_EXCESS + Per-Grid REQ", level=2)
C(doc,
    "ART_EXCESS = MAX(0, STK_TTL - 2.0 x OPT_MBQ)   [skip MIX]\n"
    "EXCESS_STK = ART_EXCESS (visible copy)\n\n"
    "Per grid: {PREFIX}_REQ = MAX(0, MBQ - (STK_TTL - agg_excess))\n"
    "ART_EXCESS column dropped after aggregation.")

# ══════════════════════════════════════════════════════════════════
# 6. STORE RANKING
# ══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("6. Phase 3: Store Ranking (ARS_STORE_RANKING)", level=1)
C(doc,
    "Per (MAJ_CAT, WERKS) — excludes MIX rows:\n"
    "  MJ_REQ    = MAX(MJ_REQ)\n"
    "  FILL_RATE = MJ_STK_TTL / MJ_MBQ\n\n"
    "  REQ_RANK  = ROW_NUMBER() ORDER BY MJ_REQ ASC\n"
    "  FILL_RANK = ROW_NUMBER() ORDER BY FILL_RATE DESC\n\n"
    "  W_SCORE   = REQ_RANK x 0.4 + FILL_RANK x 0.6\n"
    "  ST_RANK   = ROW_NUMBER() ORDER BY W_SCORE DESC\n\n"
    "ST_RANK = 1 -> best store -> gets stock first in allocation.")

# ══════════════════════════════════════════════════════════════════
# 7. WORKING TABLE
# ══════════════════════════════════════════════════════════════════
doc.add_heading("7. Phase 4: Working Table (ARS_LISTING_WORKING)", level=1)
doc.add_heading("Filters", level=2)
C(doc,
    "SELECT <identity + calculated columns>\n"
    "INTO ARS_LISTING_WORKING\n"
    "FROM ARS_LISTING\n"
    "WHERE MSA_FNL_Q > 0\n"
    "  AND OPT_REQ_WH >= 1\n"
    "  AND (VAR_FNL_COUNT / VAR_COUNT >= 0.6)   -- size availability\n"
    "  AND LISTING = 1")

doc.add_heading("Grid Hierarchy Columns", level=2)
T(doc, ["Column", "Meaning"], [
    ["GH_{hierarchy}", "Raw hierarchy match (0/1) — does this OPT belong to that grid?"],
    ["H_{hierarchy}", "Refined = GH x (REQ > 0) — belongs AND has demand"],
    ["PRI_CT%", "SUM(H_Primary) / SUM(GH_Primary) x 100"],
    ["SEC_CT%", "SUM(H_Secondary) / SUM(GH_Secondary) x 100"],
    ["ALLOC_FLAG", "1 if PRI_CT% >= 100 (eligible for allocation)"],
], cw=[3, 14.5])

# ══════════════════════════════════════════════════════════════════
# 8. MULTI-LEVEL ALLOCATION
# ══════════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading("8. Phase 5: Multi-Level Allocation (ARS_ALLOC_WORKING)", level=1)
doc.add_paragraph("Source: backend/app/services/listing_allocator.py")

doc.add_heading("8.1 Create Alloc Table + Enrichment", level=2)
C(doc,
    "ARS_LISTING_WORKING (ALLOC_FLAG=1)\n"
    "  INNER JOIN ARS_MSA_VAR_ART (FNL_Q > 0)\n"
    "  ON MAJ_CAT + GEN_ART_NUMBER + CLR + RDC\n\n"
    "Columns carried: WERKS, RDC, OPT, OPT_TYPE, ST_RANK, I_ROD,\n"
    "  DPN, SAL_D, OPT_MBQ, OPT_REQ, OPT_MBQ_WH, OPT_REQ_WH,\n"
    "  VAR_ART, SZ, FNL_Q, MRP, PAK_SZ\n\n"
    "Enrichment:\n"
    "  STK_TTL: from ARS_GRID_MJ_VAR_ART (variant-level store stock)\n"
    "  CONT:    from Master_CONT_SZ (ST -> CO -> 1/COUNT(SZ) auto-fallback)\n"
    "  SZ_MBQ = OPT_MBQ x CONT\n"
    "  SZ_REQ = MAX(0, SZ_MBQ - STK_TTL)")

doc.add_heading("8.2 Pool Tracker", level=2)
C(doc,
    "#alloc_pool temp table:\n"
    "  (RDC, MAJ_CAT, GEN_ART_NUMBER, CLR, VAR_ART, SZ)\n"
    "  FNL_Q_ORIG = initial warehouse stock\n"
    "  FNL_Q_REM  = remaining stock (decremented after each OPT allocation)\n\n"
    "Pool is SHARED per RDC — all stores with same RDC compete for same stock.")

doc.add_heading("8.3 Primary Allocation (RL > TBC > TBL, I_ROD Rounds)", level=2)
C(doc,
    "FOR each opt_type IN ['RL', 'TBC', 'TBL']:\n"
    "  max_irod = MAX(I_ROD) for this opt_type\n\n"
    "  FOR round = 1 TO max_irod:\n"
    "    eligible = OPTs WHERE OPT_TYPE = opt_type\n"
    "                      AND I_ROD >= round\n"
    "                      AND SKIP_FLAG = 0\n\n"
    "    FOR each OPT in eligible:\n"
    "      -> Allocate (waterfall all stores, ST_RANK order)\n"
    "      -> Deduct from pool\n"
    "      -> Validate (size availability check)\n"
    "      -> If fails -> find break store -> zero-out later -> SKIP\n\n"
    "Key: RL depletes pool first, TBC sees what remains, TBL gets last priority.")

doc.add_heading("8.4 Per-OPT Waterfall + Validation", level=2)
doc.add_paragraph("For each OPT, all stores allocated at once using SQL window function:")
C(doc,
    "Waterfall:\n"
    "  PARTITION BY (RDC, OPT, VAR_ART, SZ)\n"
    "  ORDER BY ST_RANK ASC, WERKS\n"
    "  ALLOC_QTY += MIN(SZ_REQ, FNL_Q_REM - prev_demand)")

doc.add_heading("Validation (per OPT, after each allocation)", level=3)
T(doc, ["#", "Check", "Action if Fails"], [
    ["B1", "MSA_FNL_Q: deduct allocated qty", "(data update)"],
    ["B2", "OPT_REQ: recalc demand", "(data update)"],
    ["B3", "Size availability: COUNT(SZ with FNL_Q_REM > 0) / total SZ", "If < 60% -> SKIP OPT for ALL stores + rounds"],
    ["B4", "Pool exhausted: MSA_FNL_Q <= 0", "SKIP"],
    ["B5", "No demand: OPT_REQ_WH < 1", "SKIP"],
    ["B6", "LISTING != 1", "SKIP"],
    ["B7", "ALLOC_FLAG != 1", "SKIP"],
], cw=[1, 7, 9.5])
doc.add_paragraph("B3 is OPT-level: if ANY store's allocation causes size avail < 60%, the entire OPT is skipped.")

doc.add_heading("Break Store Logic", level=3)
C(doc,
    "After waterfall for one OPT:\n"
    "  1. Calculate cumulative pool state per store (by ST_RANK)\n"
    "  2. Find first store where size availability drops < 60%\n"
    "  3. Zero out allocations for that store and all later stores\n"
    "  4. Restore pool for zeroed-out stores\n"
    "  5. Set SKIP_FLAG = 1 for this OPT")

doc.add_heading("Worked Example", level=3)
C(doc,
    "Pool: RDC=HYD, SHOES-123-BLACK, FNL_Q: S=30, M=100, L=80\n\n"
    "RL Round 1:\n"
    "  ST001 (rank 1): S=13, M=26, L=30 -> pool: S=17, M=74, L=50\n"
    "    Validate: 3/3 SZ with stock = 100% >= 60% -> OK\n\n"
    "  ST002 (rank 2): S=13, M=26, L=30 -> pool: S=4, M=48, L=20\n"
    "    Validate: 3/3 = 100% -> OK\n\n"
    "  ST003 (rank 3): S=4(partial), M=26, L=20(partial) -> pool: S=0, M=22, L=0\n"
    "    Validate: 1/3 SZ with stock = 33% < 60% -> SKIP\n"
    "    -> OPT skipped, no more rounds for SHOES-123-BLACK")

doc.add_page_break()
doc.add_heading("8.5 Fallback (Grid Demotion)", level=2)
doc.add_paragraph("Optional (enable_fallback = true). Relaxes grid requirements to allocate more OPTs.")
C(doc,
    "Primary grids (by seq):\n"
    "  seq=1: MJ (WERKS, MAJ_CAT)              <- ALWAYS primary\n"
    "  seq=2: MJ_RNG_SEG                       <- primary\n"
    "  seq=3: MJ_CLR                           <- primary\n"
    "  seq=4: MJ_MACRO_MVGR                    <- primary\n\n"
    "Fallback Level 1: demote seq=4 -> Secondary\n"
    "  Recalculate PRI_CT%, ALLOC_FLAG\n"
    "  Find newly eligible OPTs (ALLOC_FLAG changed 0->1)\n"
    "  Run full allocation loop for those only\n\n"
    "Fallback Level 2: demote seq=3 -> Secondary\n"
    "  More OPTs eligible -> allocate\n\n"
    "Fallback Level N: until only seq=1 remains primary\n\n"
    "Grid seq=1 is ALWAYS primary, never demoted.\n"
    "MBQ increase on fallback: placeholder for future (per STR tier).")

doc.add_heading("8.6 Reflect to Working Table", level=2)
C(doc,
    "SUM(ALLOC_QTY) per OPT -> ARS_LISTING_WORKING.ALLOC_QTY\n"
    "GROUP BY WERKS, RDC, MAJ_CAT, GEN_ART_NUMBER, CLR")

# ══════════════════════════════════════════════════════════════════
# 9. ENDPOINTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading("9. API Endpoints", level=1)
T(doc, ["Endpoint", "Method", "Purpose"], [
    ["/listing/config", "GET", "Load RDCs, stores, MAJ_CATs, settings"],
    ["/listing/generate", "POST", "Execute full listing + allocation pipeline"],
    ["/listing/settings", "POST", "Save listing variables"],
    ["/listing/preview?table=listing", "GET", "Browse ARS_LISTING"],
    ["/listing/preview?table=working", "GET", "Browse ARS_LISTING_WORKING (sorted by ST_RANK)"],
    ["/listing/preview?table=alloc", "GET", "Browse ARS_ALLOC_WORKING"],
    ["/listing/store-ranking", "GET", "Browse ARS_STORE_RANKING"],
    ["/listing/summary", "GET", "Stats by RDC, MAJ_CAT, OPT_TYPE"],
    ["/listing/create-final", "POST", "Re-create working with custom filters"],
], cw=[5, 1.5, 11])

# ══════════════════════════════════════════════════════════════════
# 10. DATA FLOW
# ══════════════════════════════════════════════════════════════════
doc.add_heading("10. Data Flow Diagram", level=1)
C(doc,
    "  Master Tables              MSA Output            Grid Output\n"
    "  +--------------+        +--------------+      +----------------+\n"
    "  | ST_MASTER    |        | ARS_MSA_     |      | ARS_GRID_MJ_   |\n"
    "  | CO_MAJ_CAT   |        | GEN_ART      |      | GEN_ART        |\n"
    "  | ST_MAJ_CAT   |        | VAR_ART      |      | VAR_ART        |\n"
    "  +------+-------+        +------+-------+      +--------+-------+\n"
    "         |                       |                       |\n"
    "    +----v--------+              |                       |\n"
    "    | Pre-Grid    |              |                       |\n"
    "    | Calculations|              |                       |\n"
    "    +----+--------+              |                       |\n"
    "         |                  +----v-----------------------v--+\n"
    "    +----v--------+         |         ARS_LISTING            |\n"
    "    | ARS_CALC_   |--DPN-->|  Part 1: Grid stock (IS_NEW=0) |\n"
    "    | ST_MAJ_CAT  |         |  Part 2: MSA new   (IS_NEW=1) |\n"
    "    | ST_ART      |--MD-->  |  Part 3: Enrich + Tag + MIX   |\n"
    "    +-------------+         |  Part 4: Grids + OPT_MBQ      |\n"
    "                            +--------------+----------------+\n"
    "                                           |\n"
    "                    +----------------------+-------------------+\n"
    "                    |                      |                   |\n"
    "            +-------v-------+    +---------v------+   +-------v--------+\n"
    "            | ARS_STORE_    |    | ARS_LISTING_   |   | ARS_ALLOC_     |\n"
    "            | RANKING       |--> | WORKING        |-->| WORKING        |\n"
    "            +---------------+    +----------------+   | (Multi-Level)  |\n"
    "                                                      | RL > TBC > TBL|\n"
    "                                                      | I_ROD rounds  |\n"
    "                                                      | + Fallback    |\n"
    "                                                      +---------------+")

# ══════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ARS_Listing_Process_SOP.docx")
doc.save(out)
print(f"Saved: {out}")
