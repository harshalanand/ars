"""
Generate ARS Listing Process Explained v2 DOCX with colored diagrams and tables.
Includes: ACS_D, ALC_D, STR, CLR capping, Focus priority, Final OPT_TYPE,
          differential size break (RL/TBC vs TBL), STR-based fallback boost, audit trail.
Run: python generate_listing_doc.py
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Colors ─────────────────────────────────────────────────────────
C_BLUE   = RGBColor(0x1F, 0x77, 0xB4); C_GREEN  = RGBColor(0x2C, 0xA0, 0x2C)
C_RED    = RGBColor(0xD6, 0x27, 0x28); C_ORANGE = RGBColor(0xFF, 0x7F, 0x0E)
C_PURPLE = RGBColor(0x94, 0x67, 0xBD); C_TEAL   = RGBColor(0x17, 0xBE, 0xCF)
C_DARK   = RGBColor(0x2C, 0x3E, 0x50); C_GREY   = RGBColor(0x7F, 0x8C, 0x8D)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

BG_BLUE_H = "1F77B4"; BG_GL = "D5F5E3"; BG_BL = "D6EAF8"; BG_OL = "FDEBD0"
BG_RL = "FADBD8"; BG_PL = "E8DAEF"; BG_TL = "D1F2EB"; BG_GR = "F2F3F4"
BG_YL = "FEF9E7"; BG_W = "FFFFFF"


def bg(cell, hx):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hx}"/>'))

def ct(cell, txt, bold=False, color=None, sz=9, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align: p.alignment = align
    r = p.add_run(txt); r.font.size = Pt(sz); r.font.name = "Calibri"
    if bold: r.bold = True
    if color: r.font.color.rgb = color

def hdr(t, headers, hbg=BG_BLUE_H):
    for i, h in enumerate(headers):
        bg(t.rows[0].cells[i], hbg); ct(t.rows[0].cells[i], h, True, C_WHITE, 9)

def rows(t, data, b1=BG_W, b2=BG_GR, bc=None):
    bc = bc or []
    for ri, rd in enumerate(data):
        row = t.add_row(); b = b1 if ri % 2 == 0 else b2
        for ci, v in enumerate(rd):
            bg(row.cells[ci], b); ct(row.cells[ci], str(v), ci in bc, sz=9)

def heading(d, txt, lv=1, c=None):
    h = d.add_heading(txt, level=lv)
    if c:
        for r in h.runs: r.font.color.rgb = c

def para(d, txt, bold=False, color=None, sz=11, sa=6):
    p = d.add_paragraph(); r = p.add_run(txt)
    r.font.size = Pt(sz); r.font.name = "Calibri"
    if bold: r.bold = True
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(sa)

def box(d, title, body, tc=C_BLUE, bx=BG_BL):
    t = d.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]; bg(c, bx)
    p = c.paragraphs[0]
    r1 = p.add_run(title + "\n"); r1.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = tc; r1.font.name = "Calibri"
    r2 = p.add_run(body); r2.font.size = Pt(10); r2.font.name = "Calibri"
    r2.font.color.rgb = C_DARK
    d.add_paragraph()

def flow(d, steps):
    for i, (txt, bh, fg) in enumerate(steps):
        t = d.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
        c = t.rows[0].cells[0]; c.width = Inches(5.5); bg(c, bh)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = fg; r.font.name = "Calibri"
        if i < len(steps) - 1:
            ap = d.add_paragraph(); ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ar = ap.add_run("\u25bc"); ar.font.size = Pt(16); ar.font.color.rgb = C_GREY

# ═══════════════════════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5); s.left_margin = Cm(2); s.right_margin = Cm(2)
doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10)

# ── TITLE PAGE ─────────────────────────────────────────────────────
doc.add_paragraph(); doc.add_paragraph()
tt = doc.add_table(rows=1, cols=1); tt.alignment = WD_TABLE_ALIGNMENT.CENTER
tc = tt.rows[0].cells[0]; bg(tc, "1A5276"); p = tc.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n ARS LISTING PROCESS v2 \n Complete Guide with Examples \n")
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = C_WHITE; r.font.name = "Calibri"
r2 = p.add_run("\nV2 Retail Auto Replenishment System\n320+ Stores | 242 MAJ_CATs\nWith ACS_D, ALC_D, STR, Color Capping, Focus Priority, Audit Trail\n\n")
r2.font.size = Pt(12); r2.font.color.rgb = RGBColor(0xAE, 0xD6, 0xF1); r2.font.name = "Calibri"
para(doc, ""); para(doc, "Purpose: Explains the entire ARS Listing Process with worked examples.", sz=11, color=C_DARK)
para(doc, "Document Date: April 2026 | Version: 2.0", sz=10, color=C_GREY)
doc.add_page_break()

# ── TOC ────────────────────────────────────────────────────────────
heading(doc, "Table of Contents", 1, C_DARK)
for item in [
    "1.  What is the Listing Process?",
    "2.  The Big Picture (9 Phases)",
    "3.  Phase 1-2: Collect Stock + STR + Add MSA",
    "4.  Phase 3: Enrich (ACS_D, ALC_D, OPT_TYPE)",
    "5.  Phase 4: Color Capping + Focus Priority",
    "6.  Phase 5: Calculate Requirements (OPT_MBQ)",
    "7.  Phase 6: Store Ranking",
    "8.  Phase 7: Working Table & Grid Coverage",
    "9.  Phase 8: Allocation (RL/TBC vs TBL rules)",
    "10. Phase 9: Final OPT_TYPE + Audit Trail",
    "11. STR-Based Fallback Boost",
    "12. Configurable Variables",
    "13. Glossary",
    "14. End-to-End Trace",
]: para(doc, item, sz=11, color=C_BLUE)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 1. What is the Listing Process
# ═══════════════════════════════════════════════════════════════════
heading(doc, "1. What is the Listing Process?", 1, C_DARK)
box(doc, "The Core Question:", '"Which products should we send to which stores, and how many pieces of each size?"', C_BLUE, BG_BL)
para(doc, "V2 Retail has 320+ stores and thousands of products. The Listing Process automates which products go where, replacing a 20-machine Excel process.", sz=10)
box(doc, "Running Example - Store DELHI-101:", "Store: DELHI-101 (tagged to WH-NORTH)\nMAJ_CAT: LEGGING\nProduct 1: 50001-BLUE (existing, 30 days old)\nProduct 2: 50002-RED (new from MSA, 5 days old)\nProduct 3: 50001-BLACK (existing, low stock - TBC)", C_GREEN, BG_GL)

# ═══════════════════════════════════════════════════════════════════
# 2. Big Picture
# ═══════════════════════════════════════════════════════════════════
heading(doc, "2. The Big Picture (9 Phases)", 1, C_DARK)
flow(doc, [
    ("Phase 1-2: Collect Stock (STK_TTL + STR) + MSA Recommendations", "2980B9", C_WHITE),
    ("Phase 3: Enrich (ACS_D, ALC_D, OPT_TYPE: MIX/RL/TBC/TBL)", "27AE60", C_WHITE),
    ("Phase 4: Color Capping (CLR_MIN/MAX) + Focus Priority (W_CAP/WO_CAP)", "16A085", C_WHITE),
    ("Phase 5: Calculate Requirements (OPT_MBQ, OPT_REQ)", "8E44AD", C_WHITE),
    ("Phase 6: Store Ranking (ST_RANK per MAJ_CAT)", "E67E22", C_WHITE),
    ("Phase 7: Working Table + Grid Coverage (ALLOC_FLAG)", "2C3E50", C_WHITE),
    ("Phase 8: Allocation  RL > TBC > TBL  (different size-break rules!)", "C0392B", C_WHITE),
    ("Phase 9: Final OPT_TYPE (RL/NL/MIX) + Full Audit Trail", "1ABC9C", C_WHITE),
])
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 3. Phase 1-2: Stock + STR + MSA
# ═══════════════════════════════════════════════════════════════════
heading(doc, "3. Phase 1-2: Collect Stock + STR + Add MSA", 1, C_DARK)
heading(doc, "Phase 1: Grid Data (Stock + Sales)", 2, C_BLUE)
para(doc, "Two key totals from the grid:", sz=10, bold=True)
t = doc.add_table(rows=1, cols=9); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["WERKS", "MAJ_CAT", "GEN_ART", "CLR", "SLOC_1000", "SLOC_2000", "L-7 SALE", "STK_TTL", "STR"])
rows(t, [
    ["DELHI-101", "LEGGING", "50001", "BLUE", "8", "2", "14", "10", "14"],
    ["DELHI-101", "LEGGING", "50001", "BLACK", "2", "1", "7", "3", "7"],
], bc=[7, 8])
doc.add_paragraph()
box(doc, "STK_TTL vs STR:", "STK_TTL = Sum of SLOC stock columns ONLY (8+2=10)\nSTR = Sum of L-7 sale columns from the SAME grid (14 pieces sold in 7 days)\n\nSTK_TTL = what's on the shelf\nSTR = how fast it's selling (used in fallback to boost MBQ)", C_BLUE, BG_BL)

heading(doc, "Phase 2: MSA Missing Options", 2, C_GREEN)
para(doc, "New products from warehouse added with STK_TTL=0, STR=0, IS_NEW=1", sz=10)
t = doc.add_table(rows=1, cols=7); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["WERKS", "MAJ_CAT", "GEN_ART", "CLR", "STK_TTL", "STR", "IS_NEW"], hbg="27AE60")
rows(t, [["DELHI-101", "LEGGING", "50002", "RED", "0", "0", "1"]], bc=[4, 5, 6])
doc.add_paragraph()
box(doc, "RDC Modes:", "All: All stores see all warehouse options\nOwn RDC: Store only sees its tagged warehouse\nCross RDC: Take FROM one warehouse, send TO another's stores", C_TEAL, BG_TL)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 4. Phase 3: Enrich + OPT_TYPE
# ═══════════════════════════════════════════════════════════════════
heading(doc, "4. Phase 3: Enrich Data & Classify Options", 1, C_DARK)
heading(doc, "Step 3.5: Add ACS_D and ALC_D", 2, C_GREEN)
t = doc.add_table(rows=1, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["WERKS", "MAJ_CAT", "ACS_D", "ALC_D"], hbg="27AE60")
rows(t, [["DELHI-101", "LEGGING", "5", "14"]])
doc.add_paragraph()
box(doc, "What do ACS_D and ALC_D mean?",
    "ACS_D (Accessories Density) = 5 means store should always have 5 pieces as base display stock\n"
    "ALC_D (Allocation Days) = 14 means send 14 days worth of additional supply for sales\n\n"
    "Formula: OPT_MBQ = ACS_D + (daily_sale_rate x ALC_D)", C_GREEN, BG_GL)

heading(doc, "Step 3.6: Initial OPT_TYPE Classification", 2, C_RED)
box(doc, "Every row gets an initial tag (may change in Phase 9):",
    "This determines allocation priority order.", C_RED, BG_RL)

# Decision tree
dt = doc.add_table(rows=1, cols=1); dt.alignment = WD_TABLE_ALIGNMENT.CENTER
bg(dt.rows[0].cells[0], "2C3E50")
ct(dt.rows[0].cells[0], "Does the store have enough stock?  (STK_TTL >= 60% x ACS_D?)", True, C_WHITE, 10, WD_ALIGN_PARAGRAPH.CENTER)
dt2 = doc.add_table(rows=1, cols=2); dt2.alignment = WD_TABLE_ALIGNMENT.CENTER
bg(dt2.rows[0].cells[0], "27AE60"); bg(dt2.rows[0].cells[1], "E74C3C")
ct(dt2.rows[0].cells[0], "YES --> RL (Repeated Listed)\nAdequate stock, replenish", True, C_WHITE, 10, WD_ALIGN_PARAGRAPH.CENTER)
ct(dt2.rows[0].cells[1], "NO --> Check further...", True, C_WHITE, 10, WD_ALIGN_PARAGRAPH.CENTER)
dt3 = doc.add_table(rows=1, cols=1); dt3.alignment = WD_TABLE_ALIGNMENT.CENTER
bg(dt3.rows[0].cells[0], "8E44AD")
ct(dt3.rows[0].cells[0], "Does warehouse have MSA stock?  (MSA_FNL_Q > 0?)", True, C_WHITE, 10, WD_ALIGN_PARAGRAPH.CENTER)
dt4 = doc.add_table(rows=1, cols=3); dt4.alignment = WD_TABLE_ALIGNMENT.CENTER
bg(dt4.rows[0].cells[0], "2980B9"); bg(dt4.rows[0].cells[1], "1ABC9C"); bg(dt4.rows[0].cells[2], "E67E22")
ct(dt4.rows[0].cells[0], "YES + STK > 0\nTBC (To Be Check)", True, C_WHITE, 10, WD_ALIGN_PARAGRAPH.CENTER)
ct(dt4.rows[0].cells[1], "YES + STK = 0\nTBL (To Be Listed)", True, C_WHITE, 10, WD_ALIGN_PARAGRAPH.CENTER)
ct(dt4.rows[0].cells[2], "NO MSA\nMIX (Excluded)", True, C_WHITE, 10, WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# Summary
t = doc.add_table(rows=1, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Init OPT_TYPE", "Condition", "Meaning", "Priority"])
for ri, (rd, bgi) in enumerate(zip([
    ["RL", "STK >= 60% x ACS_D", "Already listed, adequate stock", "1st"],
    ["TBC", "0 < STK < 60% x ACS_D, MSA > 0", "Listed, low stock, WH check", "2nd"],
    ["TBL", "STK = 0, MSA > 0", "Not in store, WH has it", "3rd"],
    ["MIX", "No MSA or poor color fill", "Excluded from allocation", "N/A"],
], [BG_GL, BG_BL, BG_TL, BG_OL])):
    row = t.add_row()
    for ci, v in enumerate(rd): bg(row.cells[ci], bgi); ct(row.cells[ci], v, ci == 0, sz=9)
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 5. Phase 4: Color Capping + Focus
# ═══════════════════════════════════════════════════════════════════
heading(doc, "5. Phase 4: Color Capping + Focus Priority", 1, C_DARK)

heading(doc, "CLR_MIN / CLR_MAX (Color Capping)", 2, C_TEAL)
box(doc, "Problem: Too many colors per GEN_ART!",
    "Example: GEN_ART 50001 (Legging) has 15 colors.\nWithout capping, store gets all 15 — too many for display.\n\n"
    "Solution:\n  CLR_MIN = 3 --> Normal: allocate top 3 colors\n  CLR_MAX = 8 --> Fallback: allow up to 8 colors\n  Colors 9-15 --> Never allocated", C_TEAL, BG_TL)

t = doc.add_table(rows=1, cols=5); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["GEN_ART", "Total Colors", "CLR_MIN", "CLR_MAX", "Behavior"], hbg="16A085")
rows(t, [
    ["50001", "15", "3", "8", "Normal: BLUE,BLACK,RED | Fallback: +GREEN,PINK... | Never: 9-15"],
], b1=BG_TL)
doc.add_paragraph()

heading(doc, "FOCUS_W_CAP / FOCUS_WO_CAP (Focus Priority)", 2, C_PURPLE)
box(doc, "Two types of focus override:",
    "FOCUS_W_CAP (Focus With Capping) = 1:\n"
    "  - TOP PRIORITY in allocation queue (RL, TBC, or TBL)\n"
    "  - Jumps ahead of non-focus options\n"
    "  - Still checks: OPT_REQ > 0 (store must need it)\n"
    "  - Still respects CLR_MIN/MAX\n"
    "  - Use case: Push specific product, but only if store needs it\n\n"
    "FOCUS_WO_CAP (Focus Without Capping) = 1:\n"
    "  - FORCE-ALLOCATE — NO requirement check needed\n"
    "  - Bypasses CLR_MIN/MAX limits\n"
    "  - Just tag with store and allocate unconditionally\n"
    "  - Use case: Promo product MUST go to store regardless", C_PURPLE, BG_PL)

t = doc.add_table(rows=1, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Product", "FOCUS_W_CAP", "FOCUS_WO_CAP", "What Happens"], hbg="8E44AD")
rows(t, [
    ["50001-BLUE", "0", "0", "Normal priority, standard rules"],
    ["50003-PINK", "1", "0", "Top priority, but only if REQ > 0"],
    ["50004-PROMO", "0", "1", "Force-allocate, skip ALL checks"],
], b1=BG_PL, b2=BG_W, bc=[0, 3])
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 6. Phase 5: Calculate Requirements
# ═══════════════════════════════════════════════════════════════════
heading(doc, "6. Phase 5: Calculate Requirements", 1, C_DARK)
heading(doc, "The Core Formula: OPT_MBQ", 2, C_PURPLE)
box(doc, "OPT_MBQ = Total pieces the store should have",
    "OPT_MBQ = ACS_D + (daily_sale_rate x ALC_D)\n\n"
    "Rate depends on AGE:\n  New (AGE < 15): rate = MAX(PER_OPT_SALE, L7_daily, AUTO_GEN_ART_SALE)\n"
    "  Old (AGE >= 15): rate = MAX(L7_daily, AUTO_GEN_ART_SALE)\n\n"
    "OPT_REQ = MAX(0, OPT_MBQ - STK_TTL)\nOPT_MBQ_WH = ACS_D + rate x (ALC_D + HOLD_DAYS)  [IS_NEW=1 only]\nOPT_REQ_WH = MAX(0, OPT_MBQ_WH - STK_TTL)", C_PURPLE, BG_PL)

para(doc, "Worked Example - 50001-BLUE (AGE=30):", sz=10, bold=True, color=C_BLUE)
t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Step", "Calculation"], hbg="8E44AD")
rows(t, [
    ["L7_daily", "14 / 7 = 2.0 per day"],
    ["Rate (AGE>=15)", "MAX(2.0, 0.8) = 2.0"],
    ["OPT_MBQ", "ACS_D + rate x ALC_D = 5 + 2.0 x 14 = 33"],
    ["OPT_REQ", "MAX(0, 33 - 10) = 23 pieces"],
], b1=BG_PL, b2=BG_W, bc=[0])
doc.add_paragraph()

para(doc, "Worked Example - 50002-RED (AGE=5, new):", sz=10, bold=True, color=C_TEAL)
t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Step", "Calculation"], hbg="1ABC9C")
rows(t, [
    ["PER_OPT_SALE", "1.5/day (from grid)"],
    ["Rate (AGE<15)", "MAX(1.5, 0, 0.8) = 1.5"],
    ["OPT_MBQ", "5 + 1.5 x 14 = 26"],
    ["OPT_REQ", "MAX(0, 26 - 0) = 26 pieces"],
], b1=BG_TL, b2=BG_W, bc=[0])
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 7. Phase 6: Store Ranking
# ═══════════════════════════════════════════════════════════════════
heading(doc, "7. Phase 6: Store Ranking", 1, C_DARK)
box(doc, "Who gets served first?", "W_SCORE = REQ_RANK x 0.4 + FILL_RANK x 0.6\nHighest W_SCORE = ST_RANK 1 (served first)", C_ORANGE, BG_OL)
t = doc.add_table(rows=1, cols=6); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Store", "MJ_REQ", "FILL_RATE", "W_SCORE", "ST_RANK", "Priority"], hbg="E67E22")
rows(t, [
    ["DELHI-101", "200", "0.167", "2.2", "1", "FIRST"],
    ["MUMBAI-55", "150", "0.333", "2.0", "2", "SECOND"],
    ["PUNE-23", "80", "0.667", "1.8", "3", "THIRD"],
], b1=BG_OL, b2=BG_YL, bc=[4, 5])
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
# 8. Phase 7: Working Table
# ═══════════════════════════════════════════════════════════════════
heading(doc, "8. Phase 7: Working Table & Grid Coverage", 1, C_DARK)
t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["#", "Filter", "Purpose"])
rows(t, [
    ["1", "MSA_FNL_Q > 0", "Warehouse must have stock"],
    ["2", "OPT_REQ_WH >= 1", "Store needs at least 1 piece"],
    ["3", "VAR ratio >= 60%", "60% of sizes available"],
    ["4", "LISTING = 1", "Product approved for sale"],
    ["5", "CLR_MIN/MAX capping", "Top N colors only"],
], b1=BG_BL, b2=BG_W)
doc.add_paragraph()
box(doc, "ALLOC_FLAG (The Gate):",
    "PRI_CT% = SUM(H_Primary) / SUM(GH_Primary) x 100\n"
    "ALLOC_FLAG = 1 if PRI_CT% >= 100%\nALLOC_FLAG = 0 if < 100% (needs fallback)", C_RED, BG_RL)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 9. Phase 8: Allocation — THE MAIN EVENT
# ═══════════════════════════════════════════════════════════════════
heading(doc, "9. Phase 8: Allocation (The Main Event!)", 1, C_DARK)
box(doc, "Waterfall Model:", "Pour from bucket (warehouse) into cups (stores) by ST_RANK.\nFOCUS_WO_CAP goes first, then FOCUS_W_CAP, then normal by rank.", C_RED, BG_RL)

heading(doc, "Allocation Priority Queue", 2, C_RED)
t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Priority", "Type", "Rule"], hbg="C0392B")
rows(t, [
    ["1st", "FOCUS_WO_CAP = 1", "Force-allocate, skip all checks"],
    ["2nd", "FOCUS_W_CAP = 1", "Top priority, but REQ > 0 required"],
    ["3rd", "RL (normal)", "Replenish by ST_RANK"],
    ["4th", "TBC (normal)", "Low-stock check by ST_RANK"],
    ["5th", "TBL (normal)", "New listing by ST_RANK"],
], b1=BG_RL, b2=BG_W, bc=[0, 1])
doc.add_paragraph()

# Waterfall example
heading(doc, "Waterfall Example (Size M, Pool = 40)", 2, C_RED)
t = doc.add_table(rows=1, cols=6); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Store", "ST_RANK", "SZ_REQ", "Gets", "Pool After", "Status"], hbg="C0392B")
wf = [
    ("DELHI-101", "1", "9", "9", "31", "ALLOCATED", BG_GL),
    ("MUMBAI-55", "2", "12", "12", "19", "ALLOCATED", BG_GL),
    ("PUNE-23", "3", "8", "8", "11", "ALLOCATED", BG_GL),
    ("JAIPUR-7", "4", "15", "11", "0", "PARTIAL", BG_OL),
    ("AGRA-12", "5", "10", "0", "0", "NOT SERVED", BG_RL),
]
for d in wf:
    row = t.add_row()
    for ci in range(6): bg(row.cells[ci], d[6]); ct(row.cells[ci], d[ci], ci in [0, 3, 5], sz=9)
doc.add_paragraph()

# ── SIZE BREAK RULES (THE BIG CHANGE) ─────────────────────────────
heading(doc, "SIZE BREAK RULES — Different for RL/TBC vs TBL!", 2, C_RED)

# Side-by-side comparison
comp = doc.add_table(rows=1, cols=2); comp.alignment = WD_TABLE_ALIGNMENT.CENTER
c_left = comp.rows[0].cells[0]; c_right = comp.rows[0].cells[1]
bg(c_left, "27AE60"); bg(c_right, "C0392B")
ct(c_left, "RL + TBC: NO Size Break Check!\n\n"
   "These are ALREADY LISTED and selling.\n"
   "Even if only 1 of 5 sizes has WH stock,\n"
   "replenish that 1 size.\n\n"
   "Partial replenishment > no replenishment\n"
   "for products already on the shelf.\n\n"
   "Example: 50001-BLUE (RL)\n"
   "Only Size M has WH stock (1/5 = 20%)\n"
   "Result: STILL ALLOCATE Size M!",
   bold=False, color=C_WHITE, sz=9)
ct(c_right, "TBL: 60% Size Break Check APPLIED!\n\n"
   "This is a FIRST-TIME listing.\n"
   "New product needs proper size coverage.\n"
   "If < 60% sizes have WH stock, SKIP.\n\n"
   "Sending 1 size of a new product\n"
   "looks bad on the shelf.\n\n"
   "Example: 50002-RED (TBL)\n"
   "Only Size M has WH stock (1/5 = 20%)\n"
   "Result: SKIP! Wait for better stock.",
   bold=False, color=C_WHITE, sz=9)
doc.add_paragraph()

box(doc, "Why the difference?",
    "RL/TBC: Customer already knows this product. If Size M is available, send it!\n"
    "TBL: Customer hasn't seen this product yet. Don't introduce it with only 1 size — wait for full range.",
    C_ORANGE, BG_OL)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 10. Phase 9: Final OPT_TYPE + Audit
# ═══════════════════════════════════════════════════════════════════
heading(doc, "10. Phase 9: Final OPT_TYPE + Audit Trail", 1, C_DARK)

heading(doc, "Final OPT_TYPE Conversion", 2, C_TEAL)
para(doc, "After allocation, initial OPT_TYPE converts to final:", sz=10)

t = doc.add_table(rows=1, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Initial", "Condition", "Final OPT_TYPE", "Meaning"], hbg="1ABC9C")
conversions = [
    ("RL", "Always", "RL (Repeated Listed)", "Existing, replenished", BG_GL),
    ("TBC", "Tagged+Allocated", "RL (Repeated Listed)", "Was uncertain, now confirmed", BG_GL),
    ("TBC", "NOT tagged/allocated", "MIX", "Warehouse couldn't serve", BG_OL),
    ("TBL", "Tagged+Allocated", "NL (New Listed)", "Successfully listed!", BG_TL),
    ("TBL", "Not allocated", "TBL (stays)", "Attempted but couldn't list", BG_YL),
    ("MIX", "Always", "MIX", "Excluded, stays excluded", BG_RL),
]
for init, cond, final, meaning, bgi in conversions:
    row = t.add_row()
    for ci, v in enumerate([init, cond, final, meaning]):
        bg(row.cells[ci], bgi); ct(row.cells[ci], v, ci in [0, 2], sz=9)
doc.add_paragraph()

# Conversion example
para(doc, "Example:", sz=10, bold=True)
t = doc.add_table(rows=1, cols=5); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Product", "Initial", "Allocated?", "Final", "Reason"], hbg="2C3E50")
rows(t, [
    ["50001-BLUE", "RL", "Yes (23 pcs)", "RL", "Existing, replenished"],
    ["50001-BLACK", "TBC", "Yes (15 pcs)", "RL", "Low stock confirmed, got WH stock"],
    ["50001-GREEN", "TBC", "No", "MIX", "WH couldn't serve"],
    ["50002-RED", "TBL", "Yes (26 pcs)", "NL", "New product listed!"],
    ["50002-YELLOW", "TBL", "No (size fail)", "TBL", "Sizes unavailable"],
], b1=BG_BL, b2=BG_W, bc=[0, 3])
doc.add_paragraph()

# Audit Trail
heading(doc, "Allocation Audit Trail", 2, C_RED)
box(doc, "CRITICAL: Every allocation must be traceable!",
    "The system records data so ANY future allocation (Listing, RL, NL, TBL, CVR)\n"
    "can be organized, tracked, and evaluated.", C_RED, BG_RL)

t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Audit Field", "Records", "Purpose"], hbg="C0392B")
audit_fields = [
    ["ALLOC_TYPE", "LISTING / RL / NL / CVR", "What type of allocation?"],
    ["ALLOC_BATCH_ID", "LST_20260419_001", "Links all rows in one run"],
    ["ALLOC_TIMESTAMP", "When allocation happened", "Time tracking"],
    ["ALLOC_USER", "Who triggered it", "Accountability"],
    ["INIT_OPT_TYPE", "RL / TBC / TBL / MIX", "Initial classification"],
    ["FINAL_OPT_TYPE", "RL / NL / MIX / TBL", "What it became"],
    ["OPT_TYPE_REASON", "Why type changed", "e.g. TBC->RL: allocated from WH"],
    ["ACS_D / ALC_D", "Parameters used", "Snapshot of config"],
    ["ALLOC_ROUND", "Which I_ROD round", "Round tracking"],
    ["ALLOC_STATUS", "ALLOCATED/PARTIAL/SKIP", "Outcome"],
    ["SKIP_REASON", "E1-E7 or B3 break", "Why skipped"],
    ["ALLOC_REMARKS", "Step-by-step trail", "RL R1:QTY=15; R2:QTY=8"],
    ["FOCUS_FLAG", "W_CAP / WO_CAP / NORMAL", "Focus priority used?"],
    ["CLR_CAP_MODE", "MIN / MAX / UNCAPPED", "Color capping applied?"],
    ["STR_BOOST_PCT", "0%/10%/20%/30%/50%", "STR boost in fallback"],
    ["POOL_START", "FNL_Q_ORIG", "Pool start quantity"],
    ["POOL_END", "FNL_Q_REM", "Pool remaining after"],
]
for ri, (f, r, p) in enumerate(audit_fields):
    row = t.add_row()
    b = BG_RL if ri % 2 == 0 else BG_W
    for ci, v in enumerate([f, r, p]):
        bg(row.cells[ci], b); ct(row.cells[ci], v, ci == 0, sz=8)
doc.add_paragraph()

box(doc, "With this audit trail you can answer:",
    "- How many NL (new listings) allocated last week?\n"
    "- Which stores got FOCUS_WO_CAP allocations?\n"
    "- What STR boost was applied to fast sellers?\n"
    "- Why was AGRA-12 skipped? -> SKIP_REASON: B3:SZ_AVAIL<60%\n"
    "- Compare RL vs NL allocation efficiency\n"
    "- Full parameter snapshot for every allocation run", C_GREEN, BG_GL)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 11. STR-Based Fallback Boost
# ═══════════════════════════════════════════════════════════════════
heading(doc, "11. STR-Based Fallback Boost", 1, C_DARK)
box(doc, "During fallback, fast-selling products get an MBQ boost!",
    "STR (Sales Turn Rate) from Phase 1 tells us how fast a product sells.\n"
    "STR Days = STK_TTL / (STR / 7) = 'How many days will current stock last?'\n\n"
    "Fast sellers need more stock urgently. Slow sellers don't.", C_PURPLE, BG_PL)

t = doc.add_table(rows=1, cols=4); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["STR Days", "Boost %", "Urgency", "Example (MBQ=26)"], hbg="8E44AD")
boost_data = [
    ["< 30 days", "+50%", "CRITICAL - runs out in < 1 month!", "26 x 1.50 = 39"],
    ["< 45 days", "+30%", "LOW - running low in 1-1.5 months", "26 x 1.30 = 34"],
    ["< 60 days", "+20%", "MODERATE - 2 months of stock", "26 x 1.20 = 31"],
    ["< 90 days", "+10%", "COMFORTABLE - could use top-up", "26 x 1.10 = 29"],
    [">= 90 days", "0%", "PLENTY - no boost needed", "26 x 1.00 = 26"],
]
bgs = [BG_RL, BG_OL, BG_YL, BG_GL, BG_GR]
for ri, (rd, bgi) in enumerate(zip(boost_data, bgs)):
    row = t.add_row()
    for ci, v in enumerate(rd): bg(row.cells[ci], bgi); ct(row.cells[ci], v, ci in [0, 1], sz=9)
doc.add_paragraph()

box(doc, "Worked Example:",
    "Product 50001-BLUE at DELHI-101:\n"
    "  STK_TTL = 10, STR (L-7 sale) = 14\n"
    "  Daily sale = 14/7 = 2.0/day\n"
    "  STR Days = 10 / 2.0 = 5 days of stock left!\n\n"
    "  5 days < 30 --> Boost = +50%\n"
    "  Original OPT_MBQ = 33\n"
    "  Boosted OPT_MBQ = 33 x 1.50 = 50 pieces\n"
    "  New OPT_REQ = MAX(0, 50 - 10) = 40 (was 23, now 40!)", C_ORANGE, BG_OL)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 12. Configurable Variables
# ═══════════════════════════════════════════════════════════════════
heading(doc, "12. Configurable Variables", 1, C_DARK)
t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Variable", "Default", "Meaning"], hbg="8E44AD")
for ri, (v, d, m) in enumerate([
    ["stock_threshold_pct", "0.6 (60%)", "RL if stock >= X% of ACS_D"],
    ["excess_multiplier", "2.0", "Excess if stock > X x OPT_MBQ"],
    ["hold_days", "0", "Extra days for IS_NEW=1"],
    ["age_threshold", "15 days", "AGE < X uses PER_OPT_SALE boost"],
    ["req_weight", "0.4", "Store rank: need weight"],
    ["fill_weight", "0.6", "Store rank: fill weight"],
    ["mix_mode", "st_maj_rng", "MIX aggregation level"],
    ["enable_fallback", "false", "Grid demotion + STR boost"],
    ["clr_min", "(per GEN_ART)", "Min colors in normal alloc"],
    ["clr_max", "(per GEN_ART)", "Max colors in fallback"],
]):
    row = t.add_row(); b = BG_PL if ri % 2 == 0 else BG_W
    for ci, val in enumerate([v, d, m]): bg(row.cells[ci], b); ct(row.cells[ci], val, ci == 0, sz=9)
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════
# 13. Glossary
# ═══════════════════════════════════════════════════════════════════
heading(doc, "13. Glossary", 1, C_DARK)
t = doc.add_table(rows=1, cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Term", "Full Name", "Meaning"], hbg="2C3E50")
for ri, (term, full, meaning) in enumerate([
    ["ACS_D", "Accessories Density", "Base display stock (replaces DPN)"],
    ["ALC_D", "Allocation Days", "Days of supply to send (replaces SAL_D)"],
    ["STR", "Sales Turn Rate", "L-7 sale total from grid (selling speed)"],
    ["OPT_MBQ", "Option Min Buy Qty", "Total pieces store should have"],
    ["OPT_REQ", "Option Requirement", "MBQ - current stock"],
    ["STK_TTL", "Stock Total", "Current stock in store"],
    ["MSA_FNL_Q", "MSA Final Qty", "Warehouse stock available"],
    ["CONT", "Contribution", "Size share % (e.g. M=35%)"],
    ["I_ROD", "Rounds", "Allocation passes count"],
    ["ALLOC_FLAG", "Allocation Flag", "1=eligible, 0=fallback"],
    ["ST_RANK", "Store Rank", "Priority (1=first)"],
    ["CLR_MIN", "Color Minimum", "Min colors, normal mode"],
    ["CLR_MAX", "Color Maximum", "Max colors, fallback mode"],
    ["FOCUS_W_CAP", "Focus With Capping", "Top priority + requirement check"],
    ["FOCUS_WO_CAP", "Focus Without Cap", "Force-allocate, skip all checks"],
    ["RL", "Repeated Listed", "Final: existing replenished"],
    ["NL", "New Listed", "Final: TBL successfully allocated"],
    ["MIX", "Mixed/Excluded", "Excluded from allocation"],
]):
    row = t.add_row(); b = BG_GR if ri % 2 == 0 else BG_W
    for ci, v in enumerate([term, full, meaning]): bg(row.cells[ci], b); ct(row.cells[ci], v, ci == 0, sz=9)
doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# 14. End-to-End
# ═══════════════════════════════════════════════════════════════════
heading(doc, "14. End-to-End Trace: 50002-RED for DELHI-101", 1, C_DARK)
t = doc.add_table(rows=1, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Step", "Detail"], hbg="1A5276")
for step, detail, bgi in [
    ("1. GRID DATA", "Not in grid (IS_NEW=1), STR=0", BG_BL),
    ("2. MSA ADDS IT", "STK_TTL=0, STR=0, IS_NEW=1", BG_BL),
    ("3. ENRICH", "ACS_D=5, ALC_D=14, AGE=5, LISTING=1, I_ROD=2", BG_GL),
    ("4. MSA DATA", "MSA_FNL_Q=200, VAR_COUNT=5, VAR_FNL_COUNT=4", BG_GL),
    ("5. INIT OPT_TYPE", "STK=0, MSA>0 --> TBL (To Be Listed)", BG_OL),
    ("6. COLOR CAP", "CLR_MIN=3, this is color #2 --> within limit", BG_TL),
    ("7. FOCUS", "W_CAP=0, WO_CAP=0 --> normal priority", BG_TL),
    ("8. OPT_MBQ", "5 + 1.5 x 14 = 26 pieces", BG_PL),
    ("9. OPT_REQ", "MAX(0, 26-0) = 26 pieces", BG_PL),
    ("10. STORE RANK", "DELHI-101 = Rank 1", BG_OL),
    ("11. WORKING TABLE", "Passes all filters", BG_BL),
    ("12. ALLOC_FLAG", "PRI_CT%=100% --> ALLOC_FLAG=1", BG_BL),
    ("13. SIZE CHECK", "TBL: 4/5=80% >= 60% --> OK!", BG_GL),
    ("14. ALLOCATION", "Gets 26 pcs: S=5, M=9, L=8, XL=4", BG_RL),
    ("15. FINAL TYPE", "TBL + allocated --> NL (New Listed)!", BG_GL),
    ("16. AUDIT", "BATCH=LST_001, INIT=TBL, FINAL=NL, BOOST=0%", BG_GR),
]:
    row = t.add_row()
    bg(row.cells[0], bgi); bg(row.cells[1], bgi)
    ct(row.cells[0], step, True, sz=9); ct(row.cells[1], detail, sz=9)
doc.add_paragraph()

box(doc, "FINAL RESULT:",
    "DELHI-101 receives 26 pieces of 50002-RED (5xS, 9xM, 8xL, 4xXL)\n"
    "Final OPT_TYPE: NL (New Listed) -- successfully listed!\n"
    "Audit: ALLOC_TYPE=LISTING, BATCH=LST_20260419_001, STR_BOOST=0%, FOCUS=NORMAL",
    C_GREEN, BG_GL)

# Footer
doc.add_paragraph()
para(doc, "Document v2 -- ACS_D, ALC_D, STR, CLR capping, Focus priority, Final OPT_TYPE,", sz=9, color=C_GREY)
para(doc, "differential size break (RL/TBC vs TBL), STR-based fallback boost, audit trail.", sz=9, color=C_GREY)
para(doc, "ARS v2.0 - V2 Retail Auto Replenishment System - April 2026", sz=9, color=C_GREY)

# ── Save ───────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "ARS_Listing_Process_Explained.docx")
doc.save(out)
print(f"Saved: {out}")
print(f"Size: {os.path.getsize(out)/1024:.1f} KB")
