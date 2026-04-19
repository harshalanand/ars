"""Generate ARS MSA Stock Calculation + Grid Builder SOP as DOCX documents."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

# ── Shared helpers ──────────────────────────────────────────────────

def make_doc():
    doc = Document()
    for section in doc.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)
    for level in range(1, 4):
        hs = doc.styles[f"Heading {level}"]
        hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
        hs.font.name = "Calibri"
    return doc

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1B3A5C"/>')
        cell._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
            if ri % 2 == 1:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="EDF2F7"/>')
                cell._tc.get_or_add_tcPr().append(shading)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="F0F0F0"/>'))

def bold(p, text):
    r = p.add_run(text); r.bold = True; return r

def bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.2 + level * 0.8)
    for r in p.runs: r.font.size = Pt(10)
    return p

def title_page(doc, title, subtitle):
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title); r.font.size = Pt(28); r.font.color.rgb = RGBColor(0x1B,0x3A,0x5C); r.bold = True
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run(subtitle); r.font.size = Pt(16); r.font.color.rgb = RGBColor(0x4A,0x6E,0x8C)
    doc.add_paragraph()
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = m.add_run("V2 Retail Auto Replenishment System\nVersion 2.0 — April 2026\n\nOwner: Akash Agarwal, Director V2 Retail\nRepository: github.com/harshalanand/ars")
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x55,0x55,0x55)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════
#  DOCUMENT 1: ARS MSA Stock Calculation SOP
# ═══════════════════════════════════════════════════════════════════════

def build_msa_sop():
    doc = make_doc()
    title_page(doc, "ARS MSA Stock Calculation", "Step-by-Step SOP & Technical Review")

    # ── TOC ──
    doc.add_heading("Table of Contents", level=1)
    for item in [
        "1. Overview",
        "2. Input & Output Tables",
        "3. Terminology",
        "4. The 9-Step MSA Calculation Pipeline",
        "   4.1 Step 1: Filter by SLOC",
        "   4.2 Step 2: Numeric Safety",
        "   4.3 Step 3: Safe Default Fill",
        "   4.4 Step 4: Segment Filter",
        "   4.5 Step 4b: Category RLS Filter",
        "   4.6 Step 5: Pivot by SLOC",
        "   4.7 Step 6: Load & Pivot Pending Allocation",
        "   4.8 Step 7: Calculate Final Quantity (FNL_Q)",
        "   4.9 Step 8: Generate Color Variants (Threshold Filter)",
        "   4.10 Step 9: Generated Colors (Aggregated)",
        "5. Output Storage & Background Jobs",
        "6. Parallel Pipeline (Batch Processing)",
        "7. API Endpoints",
        "8. Data Flow Diagram",
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(1)
        for r in p.runs: r.font.size = Pt(10)
    doc.add_page_break()

    # ── 1. Overview ──
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "The MSA (Master Stock Allocation) module calculates available warehouse stock for "
        "replenishment across all RDCs (Regional Distribution Centers). It takes raw stock data "
        "from Snowflake (via ET_MSA_STK view), deducts pending allocations, and produces three "
        "output tables used by the Listing module."
    )
    doc.add_paragraph(
        "The calculation runs as a 9-step Python pipeline using Pandas DataFrames. Results are "
        "stored to SQL Server via a background job queue."
    )
    p = doc.add_paragraph()
    bold(p, "Entry Point: "); p.add_run("POST /msa/calculate")
    p = doc.add_paragraph()
    bold(p, "Source: "); p.add_run("backend/app/services/msa_service.py")

    # ── 2. Tables ──
    doc.add_heading("2. Input & Output Tables", level=1)
    doc.add_heading("Input Tables", level=2)
    add_table(doc,
        ["Table", "Description", "Key Columns"],
        [
            ["VW_ET_MSA_STK_WITH_MASTER", "Stock view (from Snowflake sync)", "ST_CD, SLOC, ARTICLE_NUMBER, GEN_ART_NUMBER, MAJ_CAT, CLR, SZ, STK_Q"],
            ["MASTER_ALC_PEND", "Pending allocations by RDC", "RDC, ARTICLE_NUMBER, MOA (mode), QTY"],
        ],
        col_widths=[5, 5.5, 7]
    )

    doc.add_heading("Output Tables", level=2)
    doc.add_paragraph("OPT = MAJ_CAT x GEN_ART_NUMBER x CLR (same definition as listing).")
    add_table(doc,
        ["Table", "Grain", "Purpose"],
        [
            ["ARS_MSA_TOTAL", "RDC x ARTICLE_NUMBER x CLR x SZ", "Base MSA — full SLOC pivot with FNL_Q per variant"],
            ["ARS_MSA_GEN_ART", "RDC x OPT", "Generated article — aggregated across sizes (used by Listing)"],
            ["ARS_MSA_VAR_ART", "RDC x OPT x ARTICLE_NUMBER x SZ", "Variant articles — filtered by threshold (used by Alloc)"],
        ],
        col_widths=[4, 5, 8.5]
    )

    doc.add_heading("Tracking Tables", level=2)
    add_table(doc,
        ["Table", "Purpose"],
        [
            ["MSA_Calculation_Sequence", "1 row per MSA run — stores date, filters, threshold, SLOCs, row counts, status"],
            ["MSA_Column_Definitions", "Schema registry — tracks dynamically created columns per table"],
            ["msa_storage_jobs", "Background job queue — job_id, status, progress, error messages"],
        ],
        col_widths=[5, 12.5]
    )

    # ── 3. Terminology ──
    doc.add_heading("3. Terminology", level=1)
    add_table(doc,
        ["Term", "Meaning"],
        [
            ["SLOC", "Stock Location code (e.g., V01, V02_FRESH, DC01, RMS, PEND_ALC)"],
            ["STK_Q", "Stock Quantity — raw quantity per SLOC"],
            ["STK_QTY", "Total Stock Quantity = SUM of all SLOC stock after pivot"],
            ["PEND_QTY", "Pending Allocation Quantity = SUM across all MOA types"],
            ["FNL_Q", "Final Quantity = MAX(STK_QTY - PEND_QTY, 0) — available for allocation"],
            ["MOA", "Mode of Allocation (e.g., PO, TRANSFER, RETURN)"],
            ["OPT", "Option = MAJ_CAT x GEN_ART_NUMBER x CLR"],
            ["RDC", "Regional Distribution Center (ST_CD renamed to RDC in output)"],
            ["Threshold", "Minimum FNL_Q per OPT group to qualify for variant output (default 25)"],
        ],
        col_widths=[2.5, 15]
    )

    # ── 4. Pipeline ──
    doc.add_page_break()
    doc.add_heading("4. The 9-Step MSA Calculation Pipeline", level=1)
    doc.add_paragraph(
        "All 9 steps execute in-memory using Pandas DataFrames within a single API call. "
        "The pipeline transforms raw stock data through filtering, normalization, pivoting, "
        "pending deduction, and aggregation."
    )

    # Step 1
    doc.add_heading("4.1 Step 1: Filter by SLOC", level=2)
    doc.add_paragraph("Filter the raw data to include only selected stock locations.")
    code(doc,
        "Input:  Full DataFrame from VW_ET_MSA_STK_WITH_MASTER\n"
        "Filter: msa = msa[msa['SLOC'].isin(selected_slocs)]\n"
        "Output: Reduced DataFrame (only rows for selected SLOCs)"
    )
    doc.add_paragraph("User selects which SLOCs to include (e.g., V01, V02_FRESH, DC01). Unselected SLOCs are excluded from all calculations.")

    # Step 2
    doc.add_heading("4.2 Step 2: Numeric Safety", level=2)
    doc.add_paragraph("Ensure STK_Q (stock quantity) is numeric. Mixed types or NULLs are coerced to 0.")
    code(doc,
        "msa['STK_Q'] = pd.to_numeric(msa['STK_Q'], errors='coerce').fillna(0)"
    )

    # Step 3
    doc.add_heading("4.3 Step 3: Safe Default Fill", level=2)
    doc.add_paragraph(
        "Replace NULL/empty dimension values with safe defaults BEFORE pivot. "
        "This prevents NULL keys in pivot operations which would cause data loss."
    )
    add_table(doc,
        ["Column", "Default Value", "Reason"],
        [
            ["CLR", "'A'", "Color must not be NULL for pivot grouping"],
            ["SZ", "'A'", "Size must not be NULL for pivot grouping"],
            ["M_VND_NM", "'NA'", "Vendor name placeholder"],
            ["M_VND_CD", "0", "Vendor code numeric default"],
            ["MACRO_MVGR", "'NA'", "Merchandise group placeholder"],
            ["MICRO_MVGR", "'NA'", "Merchandise group placeholder"],
            ["FAB", "'NA'", "Fabric placeholder"],
            ["MVGR_MATRIX", "'NA'", "Matrix placeholder"],
            ["SSN", "'NA'", "Season placeholder"],
        ],
        col_widths=[3, 3, 11.5]
    )

    # Step 4
    doc.add_heading("4.4 Step 4: Segment Filter", level=2)
    doc.add_paragraph("Keep only relevant business segments.")
    code(doc,
        "seg_filter = ['APP', 'GM']    # Apparel + General Merchandise\n"
        "msa = msa[msa['SEG'].isin(seg_filter)]"
    )
    doc.add_paragraph("Other product segments outside ARS scope are filtered out. If SEG column is missing, all rows are retained.")

    # Step 4b
    doc.add_heading("4.5 Step 4b: Category RLS Filter", level=2)
    doc.add_paragraph(
        "Apply Row-Level Security — restrict results to the user's assigned MAJ_CATs. "
        "Admin/SuperAdmin users have no restrictions."
    )
    code(doc,
        "if user has rls_categories:\n"
        "    msa = msa[msa['MAJ_CAT'].isin(user.rls_categories)]"
    )

    # Step 5
    doc.add_heading("4.6 Step 5: Pivot by SLOC (Key Step)", level=2)
    doc.add_paragraph(
        "Transform rows into wide format — one column per SLOC, one row per article combination."
    )
    code(doc,
        "Index (row keys):  All columns EXCEPT SLOC and STK_Q\n"
        "                   (ST_CD, ARTICLE_NUMBER, GEN_ART_NUMBER, MAJ_CAT,\n"
        "                    CLR, SZ, M_VND_NM, FAB, MACRO_MVGR, etc.)\n"
        "\n"
        "Columns (pivoted): Each unique SLOC value (V01, V02_FRESH, DC01, ...)\n"
        "Values:            STK_Q (aggregated by SUM)\n"
        "Fill:              0 for missing intersections\n"
        "\n"
        "After pivot:\n"
        "  STK_QTY = SUM(all SLOC columns)    -- total stock across all locations"
    )
    doc.add_paragraph("Example: If SLOCs are V01, V02, DC01, each article row gets three stock columns plus STK_QTY.")

    # Step 6
    doc.add_page_break()
    doc.add_heading("4.7 Step 6: Load & Pivot Pending Allocation", level=2)
    doc.add_paragraph("Load MASTER_ALC_PEND, pivot by MOA (Mode of Allocation), and merge onto the main data.")
    code(doc,
        "1. Load MASTER_ALC_PEND table\n"
        "\n"
        "2. Pivot:\n"
        "   Index:   RDC, ARTICLE_NUMBER\n"
        "   Columns: MOA (PO, TRANSFER, RETURN, etc.)\n"
        "   Values:  QTY\n"
        "   → PEND_QTY = SUM(all MOA columns)\n"
        "\n"
        "3. Merge (LEFT JOIN):\n"
        "   msa_pivot LEFT JOIN pend_pivot\n"
        "   ON ST_CD = RDC AND ARTICLE_NUMBER = ARTICLE_NUMBER\n"
        "   → Unmatched rows get PEND_QTY = 0"
    )
    doc.add_heading("Example", level=3)
    add_table(doc,
        ["RDC", "ARTICLE_NUMBER", "PO", "TRANSFER", "PEND_QTY"],
        [
            ["DH24", "12345", "100", "50", "150"],
            ["DH25", "12345", "75", "0", "75"],
            ["DH26", "12345", "0", "0", "0 (no pending)"],
        ],
        col_widths=[2.5, 3.5, 2.5, 2.5, 6.5]
    )

    # Step 7
    doc.add_heading("4.8 Step 7: Calculate Final Quantity (FNL_Q)", level=2)
    doc.add_paragraph("The core calculation — available stock for allocation.")
    code(doc,
        "FNL_Q = MAX(STK_QTY - PEND_QTY, 0)\n"
        "\n"
        "Where:\n"
        "  STK_QTY  = Total stock across all SLOCs (from Step 5)\n"
        "  PEND_QTY = Total pending allocations (from Step 6)\n"
        "  Floor at 0 — no negative quantities"
    )
    doc.add_heading("Example", level=3)
    add_table(doc,
        ["RDC", "ARTICLE", "STK_QTY", "PEND_QTY", "FNL_Q"],
        [
            ["DH24", "12345", "500", "150", "350"],
            ["DH25", "12345", "100", "75", "25"],
            ["DH26", "12345", "30", "50", "0 (floored)"],
        ],
        col_widths=[2.5, 3, 3, 3, 6]
    )

    # Step 8
    doc.add_heading("4.9 Step 8: Generate Color Variants (Threshold Filter)", level=2)
    doc.add_paragraph(
        "Filter to keep only OPTs (MAJ_CAT x GEN_ART_NUMBER x CLR groups) "
        "where the total FNL_Q exceeds the threshold."
    )
    code(doc,
        "Group by: [ST_CD, MAJ_CAT, GEN_ART_NUMBER, CLR]\n"
        "Sum:      FNL_Q within each group\n"
        "Filter:   Keep rows where group SUM(FNL_Q) > threshold (default 25)\n"
        "\n"
        "Output → ARS_MSA_VAR_ART\n"
        "  Contains individual variant articles (with SZ, ARTICLE_NUMBER)\n"
        "  Only OPTs with sufficient stock pass through"
    )
    doc.add_paragraph(
        "Threshold default is 25 units. OPTs below this are considered too low for viable replenishment."
    )

    # Step 9
    doc.add_heading("4.10 Step 9: Generated Colors (Aggregated)", level=2)
    doc.add_paragraph(
        "Aggregate variant articles up to OPT level by removing ARTICLE_NUMBER and SZ from the hierarchy."
    )
    code(doc,
        "From: ARS_MSA_VAR_ART (filtered variants from Step 8)\n"
        "\n"
        "Exclude from grouping: ARTICLE_NUMBER, ARTICLE_DESC, SZ\n"
        "\n"
        "Group by: All remaining columns (hierarchy dimensions)\n"
        "          ST_CD, MAJ_CAT, GEN_ART_NUMBER, CLR, M_VND_NM, FAB, ...\n"
        "\n"
        "Aggregate: SUM for all SLOC columns + MOA columns + STK_QTY + PEND_QTY + FNL_Q\n"
        "\n"
        "Output → ARS_MSA_GEN_ART\n"
        "  One row per OPT per RDC\n"
        "  FNL_Q = aggregated available stock for that generated article"
    )

    # ── 5. Storage ──
    doc.add_page_break()
    doc.add_heading("5. Output Storage & Background Jobs", level=1)
    doc.add_paragraph(
        "After calculation completes, results are stored to SQL Server via a background job queue. "
        "The API returns immediately with the sequence_id; clients poll for storage completion."
    )
    doc.add_heading("Storage Flow", level=2)
    bullet(doc, "API returns sequence_id + job_id immediately")
    bullet(doc, "Background worker picks up job from FIFO queue")
    bullet(doc, "For each of 3 tables: TRUNCATE → batch INSERT (20k rows per batch)")
    bullet(doc, "New columns auto-created if MOA types change (schema evolution)")
    bullet(doc, "ST_CD renamed to RDC in all output tables")
    bullet(doc, "Status tracked in msa_storage_jobs table")

    doc.add_heading("Job Status Lifecycle", level=2)
    code(doc,
        "pending → running → completed\n"
        "                  → failed\n"
        "pending → cancelled (if cancelled before running)"
    )

    doc.add_heading("Column Type Inference", level=2)
    bullet(doc, "BIGINT: ARTICLE_NUMBER, GEN_ART_NUMBER, MATNR")
    bullet(doc, "NVARCHAR(200): ST_CD/RDC, SLOC, MAJ_CAT, CLR, SZ, FAB, etc.")
    bullet(doc, "FLOAT: All numeric metrics (stock, prices, quantities)")
    bullet(doc, "Dynamic MOA columns added automatically from MASTER_ALC_PEND")

    # ── 6. Parallel ──
    doc.add_heading("6. Parallel Pipeline (Batch Processing)", level=1)
    doc.add_paragraph(
        "For large-scale runs, the parallel pipeline splits work by MAJ_CAT across multiple threads."
    )
    code(doc,
        "Function: run_parallel_pipeline()\n"
        "\n"
        "1. Get all distinct MAJ_CAT from the source view\n"
        "2. Split into batches (default 5 categories per batch)\n"
        "3. Submit to ThreadPoolExecutor (default 6 workers)\n"
        "4. Each worker: creates own DB engine → runs MSA 9-step → stores results\n"
        "5. Aggregate results from all workers\n"
        "\n"
        "Performance:\n"
        "  Sequential (60 categories): ~14 hours\n"
        "  Parallel (6 workers):       ~90 minutes"
    )

    # ── 7. Endpoints ──
    doc.add_heading("7. API Endpoints", level=1)
    add_table(doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/msa/columns", "GET", "Available columns and dates for filter initialization"],
            ["/msa/distinct", "GET", "Distinct values for a column (cascading filter support)"],
            ["/msa/filter", "POST", "Apply filters and load data into memory"],
            ["/msa/calculate", "POST", "Run the 9-step MSA calculation pipeline"],
            ["/msa/save", "POST", "Manually save results to database"],
            ["/msa/results/sequences", "GET", "List recent calculation sequences"],
            ["/msa/results/{sequence_id}", "GET", "Retrieve stored results by sequence"],
            ["/msa/results/{sequence_id}/summary", "GET", "Calculation metadata/parameters"],
            ["/msa/jobs/{job_id}", "GET", "Background storage job status"],
            ["/msa/load/{config_name}", "GET", "Load saved filter configuration"],
            ["/msa/config", "POST", "Save/update filter configuration"],
        ],
        col_widths=[5, 1.5, 11]
    )

    # ── 8. Flow Diagram ──
    doc.add_heading("8. Data Flow Diagram", level=1)
    code(doc,
        "  VW_ET_MSA_STK_WITH_MASTER              MASTER_ALC_PEND\n"
        "  (Snowflake sync)                        (Pending allocations)\n"
        "  ┌──────────────────┐                    ┌──────────────────┐\n"
        "  │ ST_CD, SLOC,     │                    │ RDC, ARTICLE,    │\n"
        "  │ ARTICLE, CLR, SZ │                    │ MOA, QTY         │\n"
        "  │ STK_Q, MAJ_CAT.. │                    └────────┬─────────┘\n"
        "  └────────┬─────────┘                             │\n"
        "           │                                       │\n"
        "   Step 1: Filter SLOC                    Step 6: Pivot by MOA\n"
        "   Step 2: Numeric safety                         │\n"
        "   Step 3: Fill defaults                          │\n"
        "   Step 4: Segment filter                         │\n"
        "   Step 4b: RLS filter                            │\n"
        "           │                                      │\n"
        "   Step 5: Pivot by SLOC                          │\n"
        "           │                                      │\n"
        "           └──────────────┬───────────────────────┘\n"
        "                          │\n"
        "                  Step 6: Merge (LEFT JOIN on ST_CD + ARTICLE)\n"
        "                          │\n"
        "                  Step 7: FNL_Q = MAX(STK_QTY - PEND_QTY, 0)\n"
        "                          │\n"
        "              ┌───────────┼───────────────┐\n"
        "              │           │               │\n"
        "     ┌────────▼──────┐   │      ┌────────▼────────┐\n"
        "     │ ARS_MSA_TOTAL │   │      │ Step 8: Filter  │\n"
        "     │ (full pivot)  │   │      │ FNL_Q > 25      │\n"
        "     └───────────────┘   │      └────────┬────────┘\n"
        "                         │               │\n"
        "                         │      ┌────────▼────────┐\n"
        "                         │      │ ARS_MSA_VAR_ART │\n"
        "                         │      │ (variants)      │\n"
        "                         │      └────────┬────────┘\n"
        "                         │               │\n"
        "                         │      Step 9: Aggregate\n"
        "                         │      (remove ARTICLE, SZ)\n"
        "                         │               │\n"
        "                         │      ┌────────▼────────┐\n"
        "                         │      │ ARS_MSA_GEN_ART │\n"
        "                         │      │ (per OPT)       │\n"
        "                         │      └─────────────────┘\n"
        "                         │\n"
        "                 MSA_Calculation_Sequence (metadata)"
    )

    return doc


# ═══════════════════════════════════════════════════════════════════════
#  DOCUMENT 2: ARS Grid Builder SOP
# ═══════════════════════════════════════════════════════════════════════

def build_grid_sop():
    doc = make_doc()
    title_page(doc, "ARS Grid Builder", "Step-by-Step SOP & Technical Review")

    # ── TOC ──
    doc.add_heading("Table of Contents", level=1)
    for item in [
        "1. Overview",
        "2. Key Tables",
        "3. Terminology",
        "4. Grid Configuration (ARS_GRID_BUILDER)",
        "5. Pre-Grid Calculations (grid_calculations.py)",
        "   5.1 MAJ_CAT Level (ARS_CALC_ST_MAJ_CAT)",
        "   5.2 Master Sale (MASTER_GEN_ART_SALE)",
        "   5.3 Article Level (ARS_CALC_ST_ART)",
        "6. Grid Build & Pivot Execution",
        "   6.1 Step 1: Determine Active SLOCs",
        "   6.2 Step 2: Build Pivot Hierarchy",
        "   6.3 Step 3: Create Output Table",
        "   6.4 Step 4: Execute PIVOT",
        "   6.5 Step 5: Post-Pivot Lookups",
        "   6.6 Step 6: Grid-Level Calculations (MBQ, OPT_CNT, DISP_Q)",
        "   6.7 Step 7: Primary Key & Cleanup",
        "7. Contribution Pipeline (Master_CONT_*)",
        "8. ARS_GRID_HIERARCHY Table",
        "9. Pivot-Only Mode",
        "10. API Endpoints",
        "11. Data Flow Diagram",
    ]:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(1)
        for r in p.runs: r.font.size = Pt(10)
    doc.add_page_break()

    # ── 1. Overview ──
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "The Grid Builder module creates stock pivot tables at various hierarchy levels "
        "(MAJ_CAT, RNG_SEG, CLR, MACRO_MVGR, GEN_ART, VAR_ART). Each grid takes raw "
        "store stock from ET_STORE_STOCK, pivots it by SLOC, enriches with master data "
        "calculations, and produces an output table used by the Listing module."
    )
    doc.add_paragraph(
        "Grids are defined in ARS_GRID_BUILDER and can be run individually or in parallel "
        "(ThreadPoolExecutor, max_workers=4). Pre-grid calculations must run first to populate "
        "ARS_CALC_ST_MAJ_CAT and ARS_CALC_ST_ART."
    )
    p = doc.add_paragraph()
    bold(p, "Entry Point: "); p.add_run("POST /grid-builder/run-all  or  POST /grid-builder/grids/{id}/run")
    p = doc.add_paragraph()
    bold(p, "Source: "); p.add_run("backend/app/api/v1/endpoints/grid_builder.py + backend/app/services/grid_calculations.py")

    # ── 2. Tables ──
    doc.add_heading("2. Key Tables", level=1)
    doc.add_heading("Input Tables", level=2)
    add_table(doc,
        ["Table", "Description"],
        [
            ["ET_STORE_STOCK", "Fact table — WERKS, MATNR, SLOC, PARTICULARS_VALUE (stock qty)"],
            ["vw_master_product", "Product master view — ARTICLE_NUMBER + all hierarchy attributes"],
            ["ARS_STORE_SLOC_SETTINGS", "SLOC configuration — SLOC_CD, KPI, STATUS (Active/Inactive)"],
            ["Master_ALC_INPUT_ST_MASTER", "Store master — ST_CD, INT_DAYS, PRD_DAYS, SL_CVR, LISTING"],
            ["Master_ALC_INPUT_CO_MAJ_CAT", "Company-level MAJ_CAT config (CO base for cascade)"],
            ["Master_ALC_INPUT_ST_MAJ_CAT", "Store-level MAJ_CAT overrides"],
            ["Master_CONT_*", "Contribution tables per hierarchy level (e.g., Master_CONT_RNG_SEG)"],
        ],
        col_widths=[5.5, 12]
    )
    doc.add_heading("Output Tables", level=2)
    add_table(doc,
        ["Table", "Hierarchy", "Purpose"],
        [
            ["ARS_GRID_MJ_GEN_ART", "WERKS, MAJ_CAT, GEN_ART_NUMBER", "OPT-level stock (pivot_only)"],
            ["ARS_GRID_MJ_VAR_ART", "WERKS, MAJ_CAT, VAR_ART", "Variant-level stock (pivot_only)"],
            ["ARS_GRID_MJ_RNG_SEG", "WERKS, MAJ_CAT, RNG_SEG", "Range segment grid (with MBQ/CONT)"],
            ["ARS_GRID_MJ_CLR", "WERKS, MAJ_CAT, CLR", "Color grid (with MBQ/CONT)"],
            ["ARS_GRID_MJ_MACRO_MVGR", "WERKS, MAJ_CAT, MACRO_MVGR", "Merchandise group grid"],
            ["ARS_CALC_ST_MAJ_CAT", "ST_CD, MAJ_CAT", "Pre-grid calc — DPN, SAL_D, SAL_PD, CONT, etc."],
            ["ARS_CALC_ST_ART", "ST_CD, MAJ_CAT, GEN_ART_NUMBER, CLR", "Article-level calc"],
            ["ARS_GRID_HIERARCHY", "MAJ_CAT", "Hierarchy lookup (0/1 per grid dimension per category)"],
        ],
        col_widths=[5, 5.5, 7]
    )

    # ── 3. Terminology ──
    doc.add_heading("3. Terminology", level=1)
    add_table(doc,
        ["Term", "Meaning"],
        [
            ["SLOC", "Stock Location (V01, V02_FRESH, DC01, RMS, PEND_ALC, STK, etc.)"],
            ["KPI", "Key Performance Indicator — SLOC classification (STK = stock-related)"],
            ["STK_TTL", "Total Stock = SUM of all STK-KPI SLOCs per row"],
            ["CONT", "Contribution factor (0-1) — distributes aggregated qty to finer levels"],
            ["MBQ", "Minimum Buy Quantity = (SAL_PD x BGT_GR x SAL_D + DISP_Q x DISP_GR) x CONT"],
            ["OPT_CNT", "Option Count = ROUND(DISP_Q x DISP_GR x CONT / DPN, 0)"],
            ["DPN", "Display Norm — target stock level per store x MAJ_CAT"],
            ["SAL_D", "Sale Days — coverage period (INT_DAYS + PRD_DAYS + SL_CVR)"],
            ["SAL_PD", "Per-Day Sale rate (calculated from CM/NM sale quantities)"],
            ["DISP_Q", "Display Quantity (from ARS_CALC_ST_MAJ_CAT)"],
            ["pivot_only", "Grid flag — if true, skip post-pivot calculations (article grids)"],
            ["CO / ST", "Company level / Store level (CO = default for all stores, ST = store-specific override)"],
        ],
        col_widths=[2.5, 15]
    )

    # ── 4. Grid Config ──
    doc.add_page_break()
    doc.add_heading("4. Grid Configuration (ARS_GRID_BUILDER)", level=1)
    doc.add_paragraph("Each grid is defined as a row in ARS_GRID_BUILDER with these key fields:")
    add_table(doc,
        ["Field", "Type", "Purpose"],
        [
            ["grid_name", "NVARCHAR", "Human-readable name (e.g., 'MJ_RNG_SEG')"],
            ["hierarchy_columns", "JSON", 'Ordered columns for GROUP BY (e.g., ["WERKS","MAJ_CAT","RNG_SEG"])'],
            ["kpi_filter", "NVARCHAR", "Optional SLOC KPI filter (e.g., 'STK' = stock SLOCs only)"],
            ["output_table", "NVARCHAR", "Target table name (e.g., 'ARS_GRID_MJ_RNG_SEG')"],
            ["status", "NVARCHAR", "'Active' or 'Inactive'"],
            ["seq", "INT", "Execution sequence (also used for hierarchy column order)"],
            ["pivot_only", "BIT", "1 = skip post-pivot lookups & calculations (article grids)"],
            ["grid_group", "NVARCHAR", "'Primary' / 'Secondary' / NULL (for listing hierarchy coverage)"],
            ["use_for_opt_sale", "BIT", "1 = this grid feeds PER_OPT_SALE in listing (only ONE grid)"],
            ["weightage", "FLOAT", "Priority weight (default 1.0, informational)"],
        ],
        col_widths=[3.5, 2, 12]
    )

    # ── 5. Pre-Grid Calc ──
    doc.add_heading("5. Pre-Grid Calculations", level=1)
    doc.add_paragraph(
        "Before any grid runs, calculate_per_day_sale() builds master calculation tables. "
        "These provide DPN, SAL_D, SAL_PD, CONT, DISP_Q values that grids consume."
    )

    doc.add_heading("5.1 MAJ_CAT Level (ARS_CALC_ST_MAJ_CAT)", level=2)
    add_table(doc,
        ["Step", "Function", "What It Does"],
        [
            ["1", "Create calc (CO base)", "CO_MAJ_CAT x all stores → base table. Fallback: copy ST_MAJ_CAT."],
            ["1b", "Fill CO gaps", "Insert missing (store x MAJ_CAT) combos from CO x ST_MASTER."],
            ["2", "Overlay ST values", "ST_MAJ_CAT overrides CO per-store (non-null ST wins)."],
            ["3", "Apply defaults", "LISTING→1, I_ROD→1, growth rates→1, MANUAL_MBQ→0."],
            ["4", "SAL_D", "INT_DAYS + PRD_DAYS + SL_CVR. Priority: ST > CO > ST_MASTER."],
            ["5", "SAL_PD", "Per-day sale from CM/NM sale quantities and remaining days."],
        ],
        col_widths=[1, 3.5, 13]
    )
    doc.add_heading("SAL_PD Formula", level=3)
    code(doc,
        "IF CM_REM_D = 0                    → 0\n"
        "IF CM_REM_D >= SAL_D               → CM_SAL_Q / CM_REM_D\n"
        "IF NM_REM_D = 0                    → CM_SAL_Q / CM_REM_D\n"
        "ELSE → (CM_SAL_Q + (NM_SAL_Q / NM_REM_D) x (SAL_D - CM_REM_D)) / SAL_D"
    )

    doc.add_heading("5.2 Master Sale (MASTER_GEN_ART_SALE)", level=2)
    bullet(doc, "Ensure MAJ_CAT column exists (populate from vw_master_product if missing)")
    bullet(doc, "Compute SAL_PD in-place on ~21L rows (same formula, REM_D from ARS_CALC_ST_MAJ_CAT)")
    bullet(doc, "Feeds listing's AUTO_GEN_ART_SALE with full OPT-level coverage")

    doc.add_heading("5.3 Article Level (ARS_CALC_ST_ART)", level=2)
    doc.add_paragraph("Mirrors MAJ_CAT flow at article grain:")
    add_table(doc,
        ["Step", "What It Does"],
        [
            ["A1", "CO_ART x stores → base. Fallback: ST_ART. Drops CORE, AUTO, HH_ART."],
            ["A1b", "Fill missing (store x article) combos from CO_ART x ST_MASTER."],
            ["A2", "ST_ART overrides CO values (same cascade)."],
            ["A3", "Defaults + FOCUS_W_CAP/FOCUS_WO_CAP: Y→1, else→0."],
            ["A4", "SAL_D from ARS_CALC_ST_MAJ_CAT (article has no own SL_CVR)."],
            ["A5", "SAL_PD from MASTER_GEN_ART_SALE + ST_MAJ for REM_D."],
        ],
        col_widths=[1.5, 16]
    )

    # ── 6. Grid Build ──
    doc.add_page_break()
    doc.add_heading("6. Grid Build & Pivot Execution", level=1)
    doc.add_paragraph("Function: _build_and_run_grid(engine, grid) — executes for each active grid.")

    doc.add_heading("6.1 Step 1: Determine Active SLOCs", level=2)
    bullet(doc, "Read ARS_STORE_SLOC_SETTINGS where STATUS = 'ACTIVE'")
    bullet(doc, "Optional KPI filter (e.g., kpi_filter = 'STK' → only stock SLOCs)")
    bullet(doc, "Include PEND_ALC SLOC if active (pending allocation location)")
    bullet(doc, "Example SLOCs: RMS, STK, PEND_ALC, V01, V02_FRESH, DC01")

    doc.add_heading("6.2 Step 2: Build Pivot Hierarchy", level=2)
    bullet(doc, "WERKS comes from ET_STORE_STOCK directly")
    bullet(doc, "All other hierarchy columns resolved via LEFT JOIN to vw_master_product")
    bullet(doc, "Numeric columns (GEN_ART_NUMBER, ARTICLE_NUMBER): ISNULL → 0")
    bullet(doc, "Text columns (CLR, RNG_SEG, FAB, etc.): ISNULL → 'NA'")

    doc.add_heading("6.3 Step 3: Create Output Table", level=2)
    bullet(doc, "DROP + CREATE with hierarchy columns + one FLOAT column per active SLOC + STK_TTL")
    bullet(doc, "Auto-adds columns for newly activated SLOCs")
    bullet(doc, "Auto-drops columns for deactivated SLOCs")

    doc.add_heading("6.4 Step 4: Execute PIVOT (SQL Server)", level=2)
    code(doc,
        "WITH Stock_CTE AS (\n"
        "    SELECT hierarchy_cols, SLOC, PARTICULARS_VALUE\n"
        "    FROM ET_STORE_STOCK STK\n"
        "    LEFT JOIN vw_master_product MP ON STK.MATNR = MP.ARTICLE_NUMBER\n"
        "    INNER JOIN ARS_STORE_SLOC_SETTINGS S ON STK.SLOC = S.SLOC_CD\n"
        "    WHERE S.STATUS = 'ACTIVE' [AND S.KPI = ?]\n"
        ")\n"
        "INSERT INTO output_table (hierarchy_cols, SLOC1, SLOC2, ..., STK_TTL)\n"
        "SELECT hierarchy_cols,\n"
        "       [SLOC1], [SLOC2], ...,\n"
        "       ISNULL([STK_SLOC_1],0) + ISNULL([STK_SLOC_2],0) + ... AS STK_TTL\n"
        "FROM Stock_CTE\n"
        "PIVOT (SUM(PARTICULARS_VALUE) FOR SLOC IN ([SLOC1],[SLOC2],...)) PVT\n"
        "GROUP BY hierarchy_cols"
    )
    doc.add_paragraph("STK_TTL = sum of only STK-KPI flagged SLOCs (not all SLOCs).")

    doc.add_heading("6.5 Step 5: Post-Pivot Lookups (skip if pivot_only)", level=2)
    add_table(doc,
        ["Lookup", "Source Table", "Join", "Columns Added"],
        [
            ["1. LISTING filter", "Master_ALC_INPUT_ST_MASTER", "WERKS = ST_CD", "LISTING (then DELETE where ≠ 1)"],
            ["2. MAJ_CAT enrichment", "ARS_CALC_ST_MAJ_CAT", "WERKS = ST_CD, MAJ_CAT", "DISP_Q, DPN, SAL_D, SAL_PD, CONT, growth rates, MANUAL_MBQ"],
            ["3. Contribution", "Master_CONT_{HIER_LAST}", "WERKS + MAJ_CAT + last hier col", "CONT (with CO fallback if store-level NULL)"],
        ],
        col_widths=[2.5, 4.5, 4, 6.5]
    )
    doc.add_paragraph(
        "The {HIER_LAST} template resolves to the last hierarchy column. "
        "E.g., grid [WERKS, MAJ_CAT, RNG_SEG] → looks up Master_CONT_RNG_SEG."
    )

    doc.add_heading("6.6 Step 6: Grid-Level Calculations (skip if pivot_only)", level=2)
    doc.add_paragraph("Three KPIs are calculated using enriched data:")
    code(doc,
        "1. MBQ (Minimum Buy Quantity):\n"
        "   Raw = (SAL_PD x BGT_SL_GR_DGR) x SAL_D + (DISP_Q x DISP_GR_DGR)\n"
        "   MBQ = ROUND(Raw x CONT, 0)\n"
        "   [Growth rates default to 1 if 0/NULL; CONT=0 → MBQ=0]\n"
        "\n"
        "2. OPT_CNT (Option Count):\n"
        "   OPT_CNT = ROUND(DISP_Q x DISP_GR_DGR x CONT / DPN, 0)\n"
        "   [CONT=0 or DPN=0 → OPT_CNT=0]\n"
        "\n"
        "3. DISP_Q (effective, replaces raw value):\n"
        "   DISP_Q = ROUND(DISP_Q x CONT, 0)\n"
        "   [Runs AFTER MBQ/OPT_CNT since those use raw DISP_Q]"
    )

    doc.add_heading("6.7 Step 7: Primary Key & Cleanup", level=2)
    bullet(doc, "Fill NULL hierarchy columns: numeric → 0, text → 'NA'")
    bullet(doc, "Delete duplicate rows (keep highest STK_TTL per hierarchy group)")
    bullet(doc, "Make hierarchy columns NOT NULL")
    bullet(doc, "Add PRIMARY KEY on hierarchy columns")
    bullet(doc, "Return final row count + warnings")

    # ── 7. CONT ──
    doc.add_page_break()
    doc.add_heading("7. Contribution Pipeline (Master_CONT_*)", level=1)
    doc.add_paragraph(
        "Contribution (CONT) distributes aggregated quantities down to finer levels. "
        "Each grid hierarchy level has its own Master_CONT table."
    )
    doc.add_heading("Pattern", level=2)
    code(doc,
        "Table name: Master_CONT_{HIERARCHY_LEVEL}\n"
        "Examples:   Master_CONT_RNG_SEG, Master_CONT_CLR,\n"
        "            Master_CONT_MACRO_MVGR, Master_CONT_SZ\n"
        "\n"
        "Schema: ST_CD, MAJ_CAT, {hierarchy_col}, CONT\n"
        "\n"
        "Cascade:\n"
        "  1. Store-level: JOIN on WERKS = ST_CD + MAJ_CAT + hierarchy_col\n"
        "  2. CO fallback:  If store CONT is NULL, use ST_CD = 'CO' row\n"
        "  3. Auto fallback: If no table exists, CONT = 1 / COUNT(group)"
    )
    doc.add_heading("Example", level=2)
    add_table(doc,
        ["ST_CD", "MAJ_CAT", "RNG_SEG", "CONT"],
        [
            ["CO", "SHOES", "CASUAL", "0.40"],
            ["CO", "SHOES", "FORMAL", "0.35"],
            ["CO", "SHOES", "SPORTS", "0.25"],
            ["ST001", "SHOES", "CASUAL", "0.50  (store override)"],
            ["ST001", "SHOES", "FORMAL", "0.30  (store override)"],
        ],
        col_widths=[2.5, 3, 3, 9]
    )
    doc.add_paragraph("ST001's SPORTS CONT is NULL → falls back to CO value 0.25.")

    # ── 8. Hierarchy ──
    doc.add_heading("8. ARS_GRID_HIERARCHY Table", level=1)
    doc.add_paragraph(
        "Auto-managed companion table derived from active grid definitions. "
        "Used by the Listing module for PRI_CT%, SEC_CT%, and ALLOC_FLAG."
    )
    bullet(doc, "Base column: MAJ_CAT (primary key)")
    bullet(doc, "Dynamic columns: one per active non-article grid, named after the LAST hierarchy column")
    bullet(doc, "Values: 0 or 1 (does this MAJ_CAT participate in that grid dimension?)")
    bullet(doc, "Skips article-level grids (GEN_ART_NUMBER, ARTICLE_NUMBER, VAR_ART)")
    bullet(doc, "Automatically rebuilt when grids are created/updated/deleted")
    bullet(doc, "Column order follows grid seq")

    # ── 9. Pivot Only ──
    doc.add_heading("9. Pivot-Only Mode", level=1)
    doc.add_paragraph("When pivot_only = 1 (article-level grids):")
    bullet(doc, "Executes PIVOT + creates output table with hierarchy + SLOC columns + STK_TTL")
    bullet(doc, "SKIPS: all post-pivot lookups (LISTING filter, ARS_CALC, Master_CONT)")
    bullet(doc, "SKIPS: MBQ, OPT_CNT, DISP_Q calculations")
    bullet(doc, "Use case: GEN_ART and VAR_ART grids that only need stock pivoted for listing")

    # ── 10. Endpoints ──
    doc.add_heading("10. API Endpoints", level=1)
    add_table(doc,
        ["Endpoint", "Method", "Purpose"],
        [
            ["/grid-builder/grids", "GET", "List all grids with run status"],
            ["/grid-builder/grids", "POST", "Create new grid definition"],
            ["/grid-builder/grids/{id}", "PUT", "Update grid config"],
            ["/grid-builder/grids/{id}", "DELETE", "Delete grid + output table"],
            ["/grid-builder/grids/{id}/run", "POST", "Run single grid"],
            ["/grid-builder/run-all", "POST", "Run all Active grids in parallel (4 workers)"],
            ["/grid-builder/reorder", "PUT", "Update sequence order"],
            ["/grid-builder/columns", "GET", "List available hierarchy columns"],
            ["/grid-builder/build-calc-tables", "POST", "Build ARS_CALC tables independently"],
            ["/grid-builder/calculation-preview", "GET", "Preview pre-grid calc with timing"],
            ["/grid-builder/hierarchy/schema", "GET", "View ARS_GRID_HIERARCHY structure"],
            ["/grid-builder/hierarchy/data", "GET", "Paginated read of hierarchy data"],
        ],
        col_widths=[5.5, 1.5, 10.5]
    )

    # ── 11. Flow Diagram ──
    doc.add_heading("11. Data Flow Diagram", level=1)
    code(doc,
        "  Master Tables                    ET_STORE_STOCK           SLOC Settings\n"
        "  ┌──────────────┐              ┌──────────────────┐     ┌──────────────┐\n"
        "  │ CO_MAJ_CAT   │              │ WERKS, MATNR,    │     │ SLOC_CD, KPI │\n"
        "  │ ST_MAJ_CAT   │              │ SLOC, VALUE      │     │ STATUS       │\n"
        "  │ ST_MASTER     │              └────────┬─────────┘     └──────┬───────┘\n"
        "  └──────┬───────┘                        │                     │\n"
        "         │                                │                     │\n"
        "    ┌────▼────────────┐            ┌──────▼─────────────────────▼──┐\n"
        "    │ Pre-Grid Calc   │            │         GRID BUILD            │\n"
        "    │ (Phase 1)       │            │  Step 1: Active SLOCs        │\n"
        "    │                 │            │  Step 2: Build hierarchy      │\n"
        "    │ ARS_CALC_       │            │  Step 3: Create output table  │\n"
        "    │ ST_MAJ_CAT      │            │  Step 4: PIVOT by SLOC       │\n"
        "    │ ST_ART          │──enrich──> │  Step 5: Post-pivot lookups  │\n"
        "    └─────────────────┘            │  Step 6: MBQ, OPT_CNT, DISP │\n"
        "                                  │  Step 7: PK & cleanup        │\n"
        "  Master_CONT_*  ──CONT──>        └──────────────┬───────────────┘\n"
        "                                                 │\n"
        "                              ┌──────────────────┼──────────────────┐\n"
        "                              │                  │                  │\n"
        "                    ┌─────────▼──────┐ ┌────────▼───────┐ ┌───────▼────────┐\n"
        "                    │ ARS_GRID_MJ_   │ │ ARS_GRID_MJ_  │ │ ARS_GRID_MJ_   │\n"
        "                    │ GEN_ART        │ │ RNG_SEG       │ │ CLR / MACRO    │\n"
        "                    │ (pivot_only)   │ │ (full calc)   │ │ (full calc)    │\n"
        "                    └────────────────┘ └───────────────┘ └────────────────┘\n"
        "                              │                  │                  │\n"
        "                              └──────────────────▼──────────────────┘\n"
        "                                                 │\n"
        "                                       Consumed by LISTING\n"
        "                                       (Part 4: Grid Joins)"
    )

    return doc


# ═══════════════════════════════════════════════════════════════════════
#  GENERATE BOTH DOCUMENTS
# ═══════════════════════════════════════════════════════════════════════

base = os.path.dirname(os.path.abspath(__file__))

msa_doc = build_msa_sop()
msa_path = os.path.join(base, "ARS_MSA_Stock_Calculation_SOP.docx")
msa_doc.save(msa_path)
print(f"MSA SOP saved to: {msa_path}")

grid_doc = build_grid_sop()
grid_path = os.path.join(base, "ARS_Grid_Builder_SOP.docx")
grid_doc.save(grid_path)
print(f"Grid Builder SOP saved to: {grid_path}")
